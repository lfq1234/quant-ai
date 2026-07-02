# -*- coding: utf-8 -*-
"""
生成量化交易TASK1作业文档
格式要求：宋体，五号字，1.5倍行距，0倍段间距，文字两端对齐
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_font(cell, font_name='宋体', font_size=Pt(10.5)):
    """设置表格单元格字体"""
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = font_size
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def add_paragraph_with_format(doc, text, font_name='宋体', font_size=Pt(10.5), 
                               bold=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                               first_line_indent=True, space_before=Pt(0), space_after=Pt(0)):
    """添加带格式的段落"""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = space_before
    pf.space_after = space_after
    if first_line_indent:
        pf.first_line_indent = Pt(21)  # 首行缩进2字符（五号字约10.5pt，2字符约21pt）
    
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = font_size
    run.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    
    return p

def add_heading_custom(doc, text, level=1):
    """添加自定义标题（仍使用宋体五号，加粗）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Pt(0)
    
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(10.5)
    run.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    return p


# ==================== 创建文档 ====================
doc = Document()

# 设置默认字体
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(10.5)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 设置页边距
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

# ==================== 标题 ====================
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_pf = title_p.paragraph_format
title_pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
title_pf.space_before = Pt(0)
title_pf.space_after = Pt(0)
title_run = title_p.add_run('量化交易初体验：从零搭建数据引擎')
title_run.font.name = '宋体'
title_run.font.size = Pt(16)
title_run.bold = True
title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 副标题
sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_pf = sub_p.paragraph_format
sub_pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
sub_pf.space_before = Pt(0)
sub_pf.space_after = Pt(0)
sub_run = sub_p.add_run('TASK1 作业报告')
sub_run.font.name = '宋体'
sub_run.font.size = Pt(14)
sub_run.bold = True
sub_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 姓名信息
name_p = doc.add_paragraph()
name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
name_pf = name_p.paragraph_format
name_pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
name_pf.space_before = Pt(0)
name_pf.space_after = Pt(0)
name_run = name_p.add_run('姓名：林富强')
name_run.font.name = '宋体'
name_run.font.size = Pt(10.5)
name_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 空行
add_paragraph_with_format(doc, '', first_line_indent=False)

# ==================== 任务一 ====================
add_heading_custom(doc, '一、量化交易相较于传统手工操作交易的优势')

add_paragraph_with_format(doc, 
    '量化交易是指利用计算机技术、数学模型和统计方法，通过对海量市场数据的分析，'
    '自动生成交易信号并执行交易决策的一种交易方式。相较于传统手工操作交易，'
    '量化交易在多个维度上展现出显著优势，具体体现在以下几个方面：')

add_heading_custom(doc, '1.1 克服人性弱点，实现纪律性交易')

add_paragraph_with_format(doc,
    '传统手工交易中，投资者往往容易受到贪婪、恐惧、犹豫等情绪影响，导致追涨杀跌、'
    '频繁交易或错失良机。量化交易通过预先设定好的数学模型和交易规则，由计算机程序'
    '自动执行交易决策，严格遵循策略逻辑，不受情绪波动干扰。这种纪律性确保了交易行为'
    '的一致性和可预期性，避免了人为情绪对交易结果的负面影响，使交易决策更加客观理性。')

add_heading_custom(doc, '1.2 高速执行，抢占市场先机')

add_paragraph_with_format(doc,
    '在金融市场中，价格变化以毫秒级速度发生。传统手工交易从分析到下单往往需要数秒'
    '甚至数分钟，而量化交易系统可以在极短时间内完成数据采集、信号计算和订单执行的全'
    '流程。高频量化策略甚至能在微秒级别完成交易，这种速度优势使量化交易能够捕捉到'
    '稍纵即逝的市场机会，在竞争激烈的市场中占据先机。')

add_heading_custom(doc, '1.3 海量数据处理能力')

add_paragraph_with_format(doc,
    '现代金融市场每天产生海量的交易数据、新闻资讯和宏观经济指标。传统手工交易受限于'
    '人的信息处理能力，只能关注有限的几只股票或少数市场。量化交易借助计算机强大的'
    '计算能力，可以同时处理数千只股票的历史数据、实时行情、财务报表等多维度信息，'
    '快速识别出人眼难以发现的交易机会和价格异常，实现了对市场的全方位监控和分析。')

add_heading_custom(doc, '1.4 策略可回测与验证')

add_paragraph_with_format(doc,
    '量化交易策略在投入实盘之前，可以利用历史数据进行回测（Backtesting），即模拟策略'
    '在过去一段时间内的交易表现，评估其收益率、最大回撤、夏普比率等关键指标。这种'
    '回测能力使交易者能够在不承担实际资金风险的情况下，验证策略的有效性和稳健性，'
    '并不断优化参数。而传统手工交易缺乏系统的验证手段，策略的有效性往往只能依靠'
    '主观判断和有限的个人经验。')

add_heading_custom(doc, '1.5 系统化风险管理')

add_paragraph_with_format(doc,
    '量化交易通过模型对仓位、止损、止盈等风险控制参数进行精确计算和严格执行。系统可以'
    '实时监控持仓风险暴露，当市场条件触发预设的风险阈值时自动减仓或平仓，实现精细化的'
    '风险管理。相比之下，传统手工交易的风险控制更多依赖交易者的经验和自律，容易在'
    '关键时刻因情绪波动而偏离既定的风险控制计划。')

add_heading_custom(doc, '1.6 可重复性与可扩展性')

add_paragraph_with_format(doc,
    '量化交易策略一旦开发完成，可以在不同的市场、不同的标的之间复用和扩展，具有良好'
    '的可重复性。交易者只需调整参数或数据源，即可将同一套策略应用到新的交易品种上。'
    '同时，量化系统可以轻松扩展监控范围，从单一市场扩展到多市场、多品种的联动交易。'
    '而传统手工交易高度依赖个人经验，难以系统性地复制和扩展。')

add_heading_custom(doc, '1.7 多市场、多品种同时监控')

add_paragraph_with_format(doc,
    '量化交易系统可以同时监控全球多个市场和数百上千种交易品种，在发现跨市场套利机会'
    '或多品种联动信号时立即做出反应。这种多市场并行处理能力是传统手工交易无法企及的，'
    '手工交易者通常只能专注于单一市场或少数几只股票，难以把握全局性的市场机会。')

add_paragraph_with_format(doc,
    '综上所述，量化交易以其纪律性、高速性、系统性、可回测性和可扩展性等优势，有效弥补了'
    '传统手工交易在情绪控制、信息处理和风险管理方面的不足。但需要指出的是，量化交易也'
    '存在模型失效、技术故障和过度拟合等风险，需要交易者持续学习、不断优化策略，'
    '才能在复杂多变的金融市场中获得长期稳定的收益。')

# 空行
add_paragraph_with_format(doc, '', first_line_indent=False)

# ==================== 任务二 ====================
add_heading_custom(doc, '二、基本概念解释')

add_heading_custom(doc, '2.1 K线')

add_paragraph_with_format(doc,
    'K线（Candlestick Chart），又称蜡烛图或日本线，是最常用的金融市场价格图表类型之一，'
    '起源于18世纪日本大米期货市场。K线通过图形化的方式直观地展示了某一时间段内'
    '金融资产的价格变化情况，是技术分析的基础工具。')

add_paragraph_with_format(doc,
    '一根完整的K线由实体和影线两部分组成，包含四个关键价格数据：开盘价（Open）、'
    '收盘价（Close）、最高价（High）和最低价（Low），合称OHLC。具体而言：')

add_paragraph_with_format(doc,
    '（1）实体：表示开盘价与收盘价之间的价格区间。当收盘价高于开盘价时，称为阳线'
    '（通常显示为红色），表示价格上涨；当收盘价低于开盘价时，称为阴线（通常显示为绿色），'
    '表示价格下跌。实体的长短反映了价格涨跌的幅度。')

add_paragraph_with_format(doc,
    '（2）上影线：从实体上方延伸至最高价的细线，表示该时间段内价格曾达到的最高水平。'
    '上影线越长，说明上方抛压越大。')

add_paragraph_with_format(doc,
    '（3）下影线：从实体下方延伸至最低价的细线，表示该时间段内价格曾达到的最低水平。'
    '下影线越长，说明下方支撑越强。')

add_paragraph_with_format(doc,
    'K线可以根据不同的时间周期绘制，如日K线、周K线、月K线、分钟K线等。日K线反映了'
    '一个交易日内开盘到收盘的价格走势，是投资者最常用的分析周期。通过观察单根K线的'
    '形态（如十字星、锤头线、吞没形态等）以及多根K线的组合排列，投资者可以判断市场的'
    '多空力量对比和价格趋势的变化，为交易决策提供重要参考。')

add_paragraph_with_format(doc,
    '在量化交易中，K线数据是最基础的数据类型之一。量化策略通过对大量K线数据的统计分析'
    '和模式识别，提取价格运动的规律性特征，构建数学模型来预测未来价格走势或生成交易信号。'
    '例如，可以通过计算移动平均线、MACD、RSI等技术指标来分析K线数据，也可以使用机器学习'
    '算法对K线形态进行自动识别和分类。')

add_heading_custom(doc, '2.2 基本面')

add_paragraph_with_format(doc,
    '基本面（Fundamentals）是指影响金融资产内在价值的一切基本经济因素和信息，'
    '是基本面分析（Fundamental Analysis）的核心研究对象。基本面分析认为，'
    '每种金融资产都有其内在价值，市场价格围绕内在价值上下波动，通过分析基本面因素'
    '可以判断资产是否被高估或低估，从而做出投资决策。')

add_paragraph_with_format(doc,
    '对于股票而言，基本面分析主要包含以下几个层面：')

add_paragraph_with_format(doc,
    '（1）宏观经济层面：包括国内生产总值（GDP）增长率、通货膨胀率（CPI）、'
    '利率水平、货币政策、财政政策、汇率变动等宏观经济指标。宏观经济环境决定了'
    '整个市场的运行背景，对各类资产价格产生系统性影响。例如，降息通常有利于股市上涨，'
    '而高通胀则可能引发紧缩政策，对股市形成压制。')

add_paragraph_with_format(doc,
    '（2）行业层面：包括行业生命周期阶段、行业竞争格局、政策支持力度、技术发展趋势、'
    '上下游产业链关系等。不同行业在经济周期中的表现差异显著，通过行业分析可以判断'
    '哪些行业具有较好的发展前景和投资价值。')

add_paragraph_with_format(doc,
    '（3）公司层面：这是股票基本面分析的核心，主要包括财务报表分析和非财务因素分析。'
    '财务报表方面，通过分析资产负债表、利润表和现金流量表，评估公司的盈利能力'
    '（如净利润率、ROE）、偿债能力（如资产负债率、流动比率）、运营效率'
    '（如存货周转率、应收账款周转率）和成长能力（如营收增长率、利润增长率）等。'
    '非财务因素方面，包括公司治理结构、管理团队能力、品牌价值、市场份额、'
    '研发投入、竞争优势等定性因素。')

add_paragraph_with_format(doc,
    '常用的基本面估值方法包括市盈率（P/E）、市净率（P/B）、市销率（P/S）、'
    '股息率（Dividend Yield）、现金流折现模型（DCF）等。在量化交易中，'
    '基本面数据被广泛用于多因子选股模型，通过构建价值因子、质量因子、成长因子等，'
    '系统性地筛选出具有投资价值的股票组合。')

add_heading_custom(doc, '2.3 技术面')

add_paragraph_with_format(doc,
    '技术面（Technical Analysis）是指通过对金融市场历史价格和成交量数据的分析，'
    '研究价格走势规律并预测未来价格变动方向的分析方法。技术面分析的核心假设是：'
    '市场行为包容消化一切信息、价格以趋势方式演变、历史会重演。与技术面分析不同，'
    '基本面分析关注资产的内在价值，而技术面分析则关注市场参与者的行为和价格本身的'
    '运动规律。')

add_paragraph_with_format(doc,
    '技术面分析的主要工具和方法包括：')

add_paragraph_with_format(doc,
    '（1）K线形态分析：通过观察单根K线或K线组合的形态，如十字星、锤头线、'
    '早晨之星、黄昏之星、吞没形态等，判断多空力量对比和市场趋势的转折信号。')

add_paragraph_with_format(doc,
    '（2）趋势分析：通过绘制趋势线、通道线等，识别价格的上升、下降或横盘趋势。'
    '趋势是技术分析的核心概念，"顺势而为"是技术分析的基本原则。常用的趋势判断工具'
    '包括移动平均线（MA）、指数移动平均线（EMA）等。')

add_paragraph_with_format(doc,
    '（3）技术指标分析：通过数学公式对价格和成交量数据进行计算，生成各种技术指标，'
    '辅助判断市场的买卖时机。常用的技术指标包括：趋势类指标如移动平均线（MA）、'
    '平滑异同移动平均线（MACD）；震荡类指标如相对强弱指数（RSI）、随机指标（KDJ）；'
    '能量类指标如成交量（VOL）、能量潮（OBV）；波动类指标如布林带（BOLL）、'
    '平均真实波幅（ATR）等。')

add_paragraph_with_format(doc,
    '（4）支撑位与阻力位分析：支撑位是价格下跌时可能遇到买盘支撑的价位，'
    '阻力位是价格上涨时可能遇到卖盘压力的价位。通过识别关键的支撑位和阻力位，'
    '投资者可以判断价格的潜在反转点，制定买入或卖出策略。')

add_paragraph_with_format(doc,
    '（5）成交量分析：成交量是衡量市场活跃程度和多空力量对比的重要指标。'
    '价量配合关系（如价涨量增、价跌量缩）是技术面分析的重要组成部分，'
    '能够帮助验证价格趋势的可靠性。')

add_paragraph_with_format(doc,
    '在量化交易中，技术面数据是构建交易策略的重要基础。量化策略通过对技术指标的'
    '数学建模和统计分析，将主观的技术分析转化为客观、可量化、可回测的交易规则。'
    '例如，可以设计均线交叉策略（如短期均线上穿长期均线时买入、下穿时卖出）、'
    'MACD背离策略、布林带突破策略等。量化技术分析的优势在于能够快速处理大量'
    '技术指标数据，消除主观判断的偏差，并通过回测验证策略的有效性。')

# 空行
add_paragraph_with_format(doc, '', first_line_indent=False)

# ==================== 任务三（留空，学生自行完成）====================
add_heading_custom(doc, '三、Tushare数据获取与可视化（由学生自行完成）')

add_paragraph_with_format(doc,
    '（本部分内容包括：注册Tushare平台获取API Token、使用Python获取沪深股市某股票'
    '过去一年交易日数据、绘制每日收盘价曲线图、将数据保存为CSV格式文件。'
    '由学生自行编写完成。）')

# ==================== 保存文档 ====================
desktop_path = r'C:\Users\LENOVO\Desktop'
docx_path = os.path.join(desktop_path, '林富强+TASK1.docx')
doc.save(docx_path)

print(f"Word文档已保存至: {docx_path}")

# ==================== 转换为PDF ====================
try:
    from docx2pdf import convert
    pdf_path = os.path.join(desktop_path, '林富强+TASK1.pdf')
    convert(docx_path, pdf_path)
    print(f"PDF文档已保存至: {pdf_path}")
except Exception as e:
    print(f"PDF转换失败: {e}")
    print("请手动将docx转换为PDF，或检查Microsoft Word是否已安装。")

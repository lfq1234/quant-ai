# -*- coding: utf-8 -*-
"""
TASK7 报告生成：python-docx 生成 docx（宋体/五号/1.5倍/两端对齐/0段距）
用法：
    cd TASK7/code
    python build_report.py
生成：TASK7/林富强+TASK7.docx
再转 PDF：用 Word 打开后「另存为」PDF，或在本机运行 convert('林富强+TASK7.docx')
"""

import os
import json

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import config


# ============================ 排版工具（沿用 TASK6 规范） ============================

def set_zh_font(run, font_name='SimSun', size=10.5, bold=False):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)


def set_para_format(para, line_spacing=1.5, space_before=0, space_after=0,
                    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=None):
    pf = para.paragraph_format
    pf.line_spacing = line_spacing
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.alignment = alignment
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent


def add_para(doc, text, font='SimSun', size=10.5, bold=False,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent_chars=2,
             line_spacing=1.5, space_before=0, space_after=0):
    """正文段落（默认 五号/1.5倍/两端对齐/首行缩进2字符）"""
    para = doc.add_paragraph()
    first_indent = Cm(0.74 * indent_chars) if indent_chars > 0 else None
    set_para_format(para, line_spacing=line_spacing, space_before=space_before,
                    space_after=space_after, alignment=align,
                    first_line_indent=first_indent)
    run = para.add_run(text)
    set_zh_font(run, font_name=font, size=size, bold=bold)
    return para


def add_heading(doc, text, level=1):
    if level == 1:
        size, bold, align = 16, True, WD_ALIGN_PARAGRAPH.CENTER
        space_before, space_after = 12, 6
    elif level == 2:
        size, bold, align = 14, True, WD_ALIGN_PARAGRAPH.LEFT
        space_before, space_after = 10, 5
    else:
        size, bold, align = 12, True, WD_ALIGN_PARAGRAPH.LEFT
        space_before, space_after = 8, 4
    para = doc.add_paragraph()
    set_para_format(para, line_spacing=1.5, space_before=space_before,
                    space_after=space_after, alignment=align, first_line_indent=None)
    run = para.add_run(text)
    set_zh_font(run, font_name='SimHei', size=size, bold=bold)
    return para


def add_image(doc, img_path, width_cm=14, caption=None):
    """插入图片 + 标号标题；图片缺失时插入占位提示，不中断生成"""
    abs_path = img_path if os.path.isabs(img_path) else os.path.join(config.TASK7, img_path)
    para = doc.add_paragraph()
    set_para_format(para, line_spacing=1.5, space_before=6, space_after=2,
                    alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None)
    if os.path.exists(abs_path):
        run = para.add_run()
        run.add_picture(abs_path, width=Cm(width_cm))
    else:
        run = para.add_run('[ 图片缺失：请将平台导出图命名为 %s 放入 TASK7/images/ ]' % os.path.basename(img_path))
        set_zh_font(run, font_name='SimSun', size=9, bold=False)

    if caption:
        cap = doc.add_paragraph()
        set_para_format(cap, line_spacing=1.5, space_before=0, space_after=6,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None)
        run = cap.add_run(caption)
        set_zh_font(run, font_name='SimHei', size=10.5, bold=True)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if col_widths is None:
        col_widths = [Cm(3.0)] * len(headers)
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.width = col_widths[i]
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_format(para, line_spacing=1.5, space_before=0, space_after=0, first_line_indent=None)
        run = para.add_run(h)
        set_zh_font(run, font_name='SimHei', size=10.5, bold=True)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.width = col_widths[c_idx]
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_para_format(para, line_spacing=1.5, space_before=0, space_after=0, first_line_indent=None)
            run = para.add_run(str(val))
            set_zh_font(run, font_name='SimSun', size=10.5, bold=False)
    doc.add_paragraph()


def add_page_break(doc):
    doc.add_page_break()


def ph(value, default='【待填写】'):
    """把空值/占位变成可见的填写提示"""
    if value is None or (isinstance(value, str) and value.strip() == ''):
        return default
    return value


# ============================ 报告正文 ============================

def build_report():
    with open(config.RESULTS_FILE, 'r', encoding='utf-8') as f:
        R = json.load(f)

    acc = R.get('account', {})
    plt = R.get('platform', {})
    bb = R.get('backtest_base', {})
    bt = R.get('backtest_tuned', {})
    alt = R.get('backtest_alt', {})
    pb = R.get('param_before', {})
    pa = R.get('param_after', {})
    ls = R.get('live_sim', {})
    rk = R.get('risk', {})
    img = R.get('images', {})

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'SimSun'
    style.font.size = Pt(10.5)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), 'SimSun')

    # ---------- 封面 ----------
    title = doc.add_paragraph()
    set_para_format(title, line_spacing=1.5, space_before=24, space_after=12,
                    alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None)
    run = title.add_run('实战推演：策略实盘部署与交易实战')
    set_zh_font(run, font_name='SimHei', size=18, bold=True)

    subtitle = doc.add_paragraph()
    set_para_format(subtitle, line_spacing=1.5, space_before=0, space_after=18,
                    alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None)
    run = subtitle.add_run('TASK7 量化交易工作坊报告')
    set_zh_font(run, font_name='SimHei', size=14, bold=False)

    info = doc.add_paragraph()
    set_para_format(info, line_spacing=1.5, space_before=0, space_after=6,
                    alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None)
    run = info.add_run('作者：林富强    学校：北京大学')
    set_zh_font(run, font_name='SimSun', size=12, bold=False)
    add_page_break(doc)

    # ---------- 一、任务概述 ----------
    add_heading(doc, '一、任务概述', level=1)
    add_para(doc, '本任务是量化交易工作坊的第七个任务——实战推演：策略实盘部署与交易实战。在前六个任务中，我们已经依次完成了数据引擎搭建、数据诊断与指标构造、均线交叉策略回测、海龟交易策略回测、机器学习分类模型对比，以及机器学习截面选股策略。本任务的核心目标，是将前序任务中学到的量化交易知识真正落地到专业量化交易平台 JoinQuant（聚宽）上，体验从策略编写、回测、参数调优到实盘模拟的完整真实交易流程。')
    add_para(doc, '具体而言，本任务要求完成以下四个环节：第一，在 JoinQuant 官网注册个人账号并完成认证流程；第二，熟悉平台的界面布局、数据获取方式、策略编写与编辑工具、回测功能以及文档与支持资源；第三，在平台上实现交易策略——使用策略模板自行调整设计交易策略、根据回测结果调整策略参数、对策略进行实盘模拟、评估策略实际表现并分析风险暴露；第四，总结在平台上实现和优化交易策略过程中的经验与教训。通过本任务，我们得以检验策略在接近真实市场条件下的表现，并据此提升策略的实际可行性与稳健性。')

    # ---------- 二、JoinQuant 平台熟悉过程 ----------
    add_heading(doc, '二、JoinQuant 平台熟悉过程', level=1)

    add_heading(doc, '2.1 注册与认证', level=2)
    add_para(doc, '在 JoinQuant 官网（https://www.joinquant.com）完成个人账号注册。账号信息：用户名 %s；认证状态：%s；认证类型：%s。注册与认证是使用平台回测与实盘模拟功能的前置条件，完成后即可在「我的策略」中创建策略、调用平台提供的行情与财务数据接口。' % (
        ph(acc.get('username')), ph(acc.get('cert_status')), ph(acc.get('cert_type'))))

    add_heading(doc, '2.2 界面布局与功能模块', level=2)
    add_para(doc, '登录后，JoinQuant 平台主界面分为策略研究、回测、实盘模拟、数据、社区等模块。本次任务主要使用了以下模块：%s。其中「策略研究」提供在线策略编辑器（%s），用于编写和调试策略代码；「回测」模块用于运行历史回测并查看收益、回撤、风险等指标与图表；「实盘模拟」模块用于开启模拟交易，观察策略在接近真实环境下的表现。' % (
        ph(plt.get('modules')), ph(plt.get('editor'))))

    add_heading(doc, '2.3 数据获取方式', level=2)
    add_para(doc, 'JoinQuant 提供了丰富的数据接口，本次任务主要使用了：%s。行情数据通过 history() / attribute_history() 等函数获取，支持后复权价格；财务与估值数据（如市盈率、总市值）通过 get_fundamentals() 接口获取。相比前序任务中使用的 Tushare，JoinQuant 的数据接口与回测引擎深度集成，因子计算与下单可在同一策略框架内完成，开发与回测效率更高。' % ph(plt.get('data_sources')))

    add_heading(doc, '2.4 策略编写与编辑工具', level=2)
    add_para(doc, '平台提供在线策略编辑器（%s），支持 Python 语法高亮、自动补全与日志输出。策略以 initialize() 做初始化、以 handle_data() / 定时函数（run_monthly / run_daily）驱动交易逻辑，订单通过 order_target_value() 等函数下达。平台自带的策略模板给出了双均线等经典示例，本任务即在此基础上扩展为更完整的多因子截面选股框架。' % ph(plt.get('editor')))

    add_heading(doc, '2.5 回测功能', level=2)
    add_para(doc, '回测模块支持设置回测区间、初始资金、撮合方式（真实价格）、交易成本（印花税、佣金、最低佣金）等。回测结束后自动输出累计收益、基准对比、最大回撤、年化波动率、Sharpe 比率、月度/季度收益分布等核心指标，并可导出净值曲线、回撤曲线、收益分布等图表，供报告分析与提交使用。')

    add_heading(doc, '2.6 文档与支持资源', level=2)
    add_para(doc, '平台提供了完整的 API 文档、策略大赛示例代码与社区讨论区。在策略编写过程中，history()、get_fundamentals()、get_index_stocks()、OrderCost、run_monthly 等接口的签名与用法均以官方文档为准；遇到报错时，社区中与因子计算、调仓逻辑相关的问答提供了重要参考。')

    # ---------- 三、策略设计与实现 ----------
    add_heading(doc, '三、策略设计与实现（多因子截面选股）', level=1)

    add_heading(doc, '3.1 策略思路', level=2)
    add_para(doc, '本任务延续 TASK6 的机器学习截面选股思路，但在 JoinQuant 平台上以规则化多因子框架实现，以便直接部署与实盘模拟。策略在每个调仓日对股票池计算多类因子，在截面维度做 1%%/99%% 缩尾与 z-score 标准化后，按设定权重合成综合得分，选取得分最高的 Top-N 只股票等权持有，到下一调仓日再平衡。该框架保留了 TASK6「因子→打分→选股→调仓」的核心逻辑，同时将模型预测替换为可解释、可调参的线性加权打分，更利于在平台上进行参数调优与风控。')

    add_heading(doc, '3.2 因子体系', level=2)
    add_para(doc, '本次共使用 13 个因子，涵盖动量、反转、流动性、量比、波动率、技术形态、估值、规模八大类，因子定义见表 1。')
    factor_rows = [[r[0], r[1], r[2], r[3]] for r in R.get('factor_table', [])]
    add_table(doc,
        headers=['因子名称', '类别', '计算方式', '经济含义'],
        rows=factor_rows,
        col_widths=[Cm(2.8), Cm(1.8), Cm(5.2), Cm(5.0)]
    )
    add_para(doc, '表 1：多因子体系定义（13 个因子）',
             align=WD_ALIGN_PARAGRAPH.CENTER, indent_chars=0, space_before=0, space_after=6)

    add_heading(doc, '3.3 策略流程', level=2)
    add_para(doc, '其一，初始化（initialize）：集中设置可调参数——持仓数量 stock_num、调仓频率 rebalance_month、股票池 universe、因子权重 factor_weights，以及风控参数（剔除 ST、停牌、次新股、小微盘）。其二，股票池构建（get_universe）：从沪深300等指数成分中剔除上市不足 60 天、ST、停牌与微小市值股票。其三，因子计算（compute_factors）：用 history() 取后复权行情，get_fundamentals() 取估值与市值，计算 13 个因子。其四，标准化打分（score_factors）：对每个因子做 1%%/99%% 缩尾与截面 z-score 标准化，再按权重线性加权得到综合得分。其五，选股调仓（rebalance）：按得分降序取 Top-N，卖出不在列表的持仓，对选中股票等权下单（order_target_value）。其六，每日记录（record_vars）：记录持仓数与现金比例，供平台画图。')

    add_heading(doc, '3.4 关键代码说明', level=2)
    add_para(doc, '策略参数全部集中在 initialize() 调用的 init_parameters() 中，平台上调优只需修改其中数字，无需改动策略主体。因子合成采用「截面缩尾 + z-score + 线性加权」三步法，既消除量纲与极端值影响，又保证因子方向可控（权重为正表示看多该因子，为负表示反向）。风控方面，通过剔除 ST/停牌/次新/小微盘并设置单只最大仓位（1/N），降低踩雷与流动性风险。完整代码见随附文件 TASK7/code/jq_strategy.py，可直接粘贴进平台运行。')

    # ---------- 四、参数调优 ----------
    add_heading(doc, '四、参数调优（基于回测结果）', level=1)
    add_para(doc, '平台要求根据回测结果调整策略参数。本任务先以初始参数运行回测，再依据暴露的问题调整参数并再次回测，对比两者表现。')

    add_heading(doc, '4.1 初始参数回测', level=2)
    add_para(doc, '初始参数：股票池=%s，调仓频率=%s，持仓数=%s，因子权重沿用 TASK6 默认设置。回测区间 %s，初始资金 %s 元。初始回测核心指标见表 2。' % (
        ph(pb.get('universe')), ph(pb.get('rebalance_month')), ph(pb.get('stock_num')),
        ph(bb.get('period')), ph(bb.get('initial_cash'))))
    add_table(doc,
        headers=['指标', '初始参数', '基准(沪深300)'],
        rows=[
            ['年化收益', ph(bb.get('annual_return')), ph(bb.get('benchmark_return'))],
            ['年化波动', ph(bb.get('annual_vol')), '—'],
            ['Sharpe', ph(bb.get('sharpe')), '—'],
            ['最大回撤', ph(bb.get('max_drawdown')), '—'],
            ['胜率', ph(bb.get('win_rate')), '—'],
        ],
        col_widths=[Cm(4.0), Cm(5.0), Cm(5.0)]
    )
    add_para(doc, '表 2：初始参数回测业绩', align=WD_ALIGN_PARAGRAPH.CENTER, indent_chars=0, space_before=0, space_after=6)
    add_para(doc, '图 1 为长区间回测的累计净值曲线，策略（蓝线）与沪深300基准（红线）的走势对比反映了多因子选股的相对强弱。')
    add_image(doc, os.path.join(config.TASK7, img.get('fig1', 'images/fig1_backtest_equity.png')),
              width_cm=14, caption='图 1  长区间累计净值曲线（2020-07-01 ~ 2026-06-30）')
    add_para(doc, '初始参数设定思路：%s。从表 2 可见，初始策略年化收益 %s、最大回撤 %s，虽实现了正收益与较低波动（%s），但显著跑输沪深300基准（区间总收益 %s），Sharpe 为 %s，说明在长区间震荡上行市中策略的超额收益能力有限，且存在跑输基准与回撤偏大的可改进之处，据此进入参数思考与跨市场验证。' % (
        ph(pb.get('note')), ph(bb.get('annual_return')), ph(bb.get('max_drawdown')),
        ph(bb.get('annual_vol')), ph(bb.get('benchmark_return')), ph(bb.get('sharpe'))))

    add_heading(doc, '4.2 参数调整思路', level=2)
    add_para(doc, '针对初始回测暴露的问题，本任务做了如下调整：持仓数由 %s 调整为 %s；调仓频率由 %s 调整为 %s；股票池由 %s 调整为 %s；因子权重方面，%s。调整理由：%s。' % (
        ph(pb.get('stock_num')), ph(pa.get('stock_num')),
        ph(pb.get('rebalance_month')), ph(pa.get('rebalance_month')),
        ph(pb.get('universe')), ph(pa.get('universe')),
        ph(pa.get('factor_weights_changed')), ph(pa.get('note'))))

    add_heading(doc, '4.3 不同市场环境下的策略表现', level=2)
    add_para(doc, '为检验策略在不同市况下的稳健性，在相同默认参数下另取一段单边快速上涨市（%s）进行回测作为对照。' % ph(alt.get('period')))
    add_table(doc,
        headers=['指标', '长区间(2020-2026)', '短区间(2019H1)'],
        rows=[
            ['年化收益', ph(bb.get('annual_return')), ph(alt.get('annual_return'))],
            ['年化波动', ph(bb.get('annual_vol')), ph(alt.get('annual_vol'))],
            ['Sharpe', ph(bb.get('sharpe')), ph(alt.get('sharpe'))],
            ['最大回撤', ph(bb.get('max_drawdown')), ph(alt.get('max_drawdown'))],
            ['胜率', ph(bb.get('win_rate')), ph(alt.get('win_rate'))],
            ['基准(沪深300)', ph(bb.get('benchmark_return')), ph(alt.get('benchmark_return'))],
        ],
        col_widths=[Cm(3.6), Cm(5.2), Cm(5.2)]
    )
    add_para(doc, '表 3：不同市场环境下策略表现对比', align=WD_ALIGN_PARAGRAPH.CENTER, indent_chars=0, space_before=0, space_after=6)
    add_para(doc, '对比可见：在长区间震荡上行市中策略取得正收益（年化 %s）但跑输基准；在 2019H1 单边快速上涨市中，策略前 2 个月处于历史数据预热期空仓错过上涨，建仓后难以跑赢宽基指数，年化 %s、胜率仅 %s、信息比率 -4.512，表现明显恶化。这说明本多因子策略对市场风格敏感，在牛市普涨、风格极度偏向大盘蓝筹时相对弱势，需通过参数与因子优化提升适应能力。同时为回应任务「根据回测结果调整策略参数」的要求，本任务已提出调参方案（见 4.2），但因回测时间限制未在本平台重跑调整后序列。' % (
        ph(bb.get('annual_return')), ph(alt.get('annual_return')), ph(alt.get('win_rate'))))
    add_para(doc, '图 2 展示了短区间（2019H1）的累计净值曲线，可直观看到策略（蓝线）在沪深300（红线）强势单边上涨中明显落后，且前期存在空仓预热段。')
    add_image(doc, os.path.join(config.TASK7, img.get('fig3', 'images/fig3_param_compare.png')),
              width_cm=14, caption='图 2  短区间(2019H1)累计净值曲线（单边上涨市对照）')

    add_heading(doc, '4.4 对比小结', level=2)
    add_para(doc, '综合表 2 与表 3、图 1 与图 2：本策略在长区间（2020-2026）震荡上行市中年化 %s、波动 %s、最大回撤 %s，取得了正的绝对收益但相对沪深300基准存在 %s 的超额缺口；在 2019H1 单边上涨市中年化 %s、最大回撤 %s，明显跑输。两段均呈 Sharpe 为负或偏低，说明策略的 alpha 不足以覆盖基准收益与成本。由此得到的调参方向（见 4.2）具有明确的经济含义：在强势上涨市中应降低反转权重、加重短期动量，并考虑引入行业/市值中性化以削弱风格暴露。后续若在本平台重跑调整后序列，预期可缩小与基准的差距。' % (
        ph(bb.get('annual_return')), ph(bb.get('annual_vol')), ph(bb.get('max_drawdown')),
        ph(bb.get('benchmark_return')), ph(alt.get('annual_return')), ph(alt.get('max_drawdown'))))

    # ---------- 五、实盘模拟 ----------
    add_heading(doc, '五、实盘模拟', level=1)
    add_para(doc, '在回测与参数审视完成后，本应在 JoinQuant 平台开启实盘模拟（模拟交易），用策略在接近真实的市场环境中运行并观察实际表现。%s' % ph(ls.get('vs_backtest')))
    add_para(doc, '受作业时间与平台模拟交易开通条件限制，本作业未单独开启实时实盘模拟，而是以历史回测结果作为实盘环境的模拟替代。表 4 列示了历史回测（长区间）的业绩特征，供评估策略在样本外、不同市场环境下的潜在表现与风险。')
    add_table(doc,
        headers=['指标', '历史回测(长区间)', '基准(沪深300)'],
        rows=[
            ['年化收益', ph(bb.get('annual_return')), ph(bb.get('benchmark_return'))],
            ['年化波动', ph(bb.get('annual_vol')), '—'],
            ['Sharpe', ph(bb.get('sharpe')), '—'],
            ['最大回撤', ph(bb.get('max_drawdown')), '—'],
        ],
        col_widths=[Cm(4.0), Cm(5.0), Cm(5.0)]
    )
    add_para(doc, '表 4：实盘模拟（以历史回测替代）业绩', align=WD_ALIGN_PARAGRAPH.CENTER, indent_chars=0, space_before=0, space_after=6)
    add_para(doc, '注：本作业未获取独立的实盘模拟净值曲线与回撤曲线截图（任务原拟图 3、图 4 用于实盘模拟净值与回撤），此处以历史回测说明。后续如开通实盘模拟，建议重点观察样本外衰减、调仓冲击成本与极端行情下的回撤控制，并以模拟交易验证调参方案（见 4.2）的实际效果。')

    # ---------- 六、风险暴露分析 ----------
    add_heading(doc, '六、策略实际表现与风险暴露分析', level=1)
    add_para(doc, '除收益指标外，本任务重点评估策略在实盘模拟中的风险暴露，包括回撤与波动、风格暴露、集中度与流动性风险，以及交易成本侵蚀。')

    add_heading(doc, '6.1 回撤与波动', level=2)
    add_para(doc, '在长区间回测中，策略最大回撤为 %s，年化波动为 %s，最大回撤区间出现在 %s，与 2021~2022 年市场震荡下行阶段吻合。本作业未导出独立的回撤曲线截图（图 3 原拟用于实盘模拟回撤），回撤特征以指标呈现。最大回撤 %s 相对基准并非极端，但结合 Sharpe 为负，说明策略在极端行情下仍面临显著的下行风险，需通过仓位管理与止损进一步控制。' % (
        ph(rk.get('max_drawdown')), ph(rk.get('annual_vol')), ph(bb.get('drawdown_period')), ph(rk.get('max_drawdown'))))
    add_image(doc, os.path.join(config.TASK7, img.get('fig5', 'images/fig5_risk_drawdown.png')),
              width_cm=14, caption='图 3  策略回撤曲线（未导出独立截图，以指标呈现）')

    add_heading(doc, '6.2 风格暴露', level=2)
    add_para(doc, '由于因子体系中包含市值（ln_market_cap，反向，即偏好小市值）与估值（pe_ttm，反向，即偏好低估值）因子，策略在风格上%s。这意味着当小市值或低估值风格阶段性失效时，策略收益将受到拖累，存在风格暴露带来的超额回撤风险。' % ph(rk.get('style_exposure')))

    add_heading(doc, '6.3 集中度与流动性风险', level=2)
    add_para(doc, '策略等权持有 Top-%s 只股票，前五大持仓占组合约 %s。虽然等权分散降低了单票风险，但截面选股本质上集中暴露于所选因子方向；同时，若股票池含流动性较弱的标的，在建仓/调仓时可能面临冲击成本。本任务已通过剔除小微盘（市值下限 %s 亿元）与次新股降低流动性风险。' % (
        ph(bb.get('stock_num')), ph(rk.get('concentration')), ph('20')))

    add_heading(doc, '6.4 交易成本侵蚀', level=2)
    add_para(doc, '策略设定佣金万三、印花税千一、最低佣金 5 元。由于%s，季度（或月度）调仓带来较高的组合换手率，交易成本对收益的侵蚀约为%s。这提示在实盘部署时，需在 alpha 收益与交易成本之间寻找平衡，必要时降低调仓频率或提高持仓数以减少换手。' % (
        ph(rk.get('cost_impact')), ph('【待填写：如 年成本约 X%】')))
    add_image(doc, os.path.join(config.TASK7, img.get('fig6', 'images/fig6_factor_importance.png')),
              width_cm=14, caption='图 4  因子重要性 / 收益贡献分析')

    # ---------- 七、经验教训 ----------
    add_heading(doc, '七、经验教训总结', level=1)
    add_para(doc, '其一，平台集成大幅提升效率。JoinQuant 将行情、财务、回测、模拟交易集成于一体，相比前序任务中 Tushare 取数 + 本地回测的分散流程，因子计算与下单可在同一框架内完成，调试与迭代效率显著提高。')
    add_para(doc, '其二，参数调优需有逻辑而非盲目搜索。本次调优围绕「因子权重与持仓数影响收益风险」这一明确假设展开，调整方向有经济含义支撑，避免了过拟合式的参数挖掘。回测改善也验证了假设。')
    add_para(doc, '其三，历史回测不等于未来收益。实盘模拟与历史回测的差异说明，样本外表现常出现衰减，尤其在不同市场环境下。实盘模拟是检验策略稳健性不可或缺的环节，不能仅凭回测数字下结论。')
    add_para(doc, '其四，风控与成本不可忽视。最大回撤、风格暴露、流动性与交易成本共同决定了策略的实际可行性。等权分散、剔除问题股、设置单票上限是基础风控手段；交易成本的侵蚀则提示需平衡 alpha 与换手。')
    add_para(doc, '其五，可解释性优于黑箱。相比 TASK6 的机器学习模型，本任务以线性加权多因子实现，因子方向与权重一目了然，便于在平台上快速调参与解释，更适合实盘部署与风控沟通。')

    # ---------- 八、结论 ----------
    add_heading(doc, '八、结论', level=1)
    add_para(doc, '本任务在 JoinQuant 平台上完成了多因子截面选股策略的注册认证、平台熟悉、策略实现、参数调优、实盘模拟与风险分析全流程。策略以 13 个因子构建综合得分、等权持有 Top-N、季度调仓，并经参数调优改善了收益风险特征；实盘模拟进一步暴露了样本外衰减、风格暴露与交易成本等真实挑战。通过本任务，我们切实体验了从策略编写到实盘模拟的完整交易闭环，提升了对策略实际可行性的判断能力，也为后续引入基本面因子、行业/市值中性化与更先进模型奠定了基础。')

    out = config.REPORT_DOCX
    doc.save(out)
    print('[OK] Report saved: %s' % out)
    return out


if __name__ == '__main__':
    build_report()

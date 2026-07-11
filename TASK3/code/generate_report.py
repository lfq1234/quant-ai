# -*- coding: utf-8 -*-
"""
报告生成模块：读取回测结果，生成 Word 文档并转换为 PDF。
运行方式: cd code && python generate_report.py

格式要求：宋体，五号字，1.5倍行距，0倍段间距，文字两端对齐
内容覆盖：入场规则、出场规则、仓位管理、风险控制、绩效指标、Python实现、参数对比与总结
"""

import os
import json

from config import BASE_DIR, DATA_DIR, IMAGE_DIR
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# ==================== 加载回测结果 ====================
results_path = os.path.join(DATA_DIR, 'backtest_results.json')
with open(results_path, 'r', encoding='utf-8') as f:
    results = json.load(f)

base = results['base']
position_results = results['positions']
period_results = results['periods']
stock_results = results['stocks']
trade_records = results['trades']

# ==================== 图片路径辅助 ====================
def img_path(subdir: str, filename: str) -> str:
    """获取图片完整路径"""
    if subdir:
        return os.path.join(IMAGE_DIR, subdir, filename)
    return os.path.join(IMAGE_DIR, filename)


# ==================== 文档格式函数 ====================
def add_paragraph(doc, text, bold=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  first_line_indent=True, font_size=Pt(10.5)):
    """添加带格式的段落（宋体五号，1.5倍行距，0段间距，两端对齐）"""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    if first_line_indent:
        pf.first_line_indent = Pt(21)
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = font_size
    run.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p


def add_heading(doc, text, font_size=Pt(10.5)):
    """添加标题（宋体五号加粗，左对齐）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Pt(0)
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = font_size
    run.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p


def add_image(doc, img_path_full, caption, width=Inches(6.0)):
    """添加图片和图注"""
    if not os.path.exists(img_path_full):
        print(f'警告：图片不存在 {img_path_full}')
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Pt(0)
    run = p.add_run()
    run.add_picture(img_path_full, width=width)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf2 = p2.paragraph_format
    pf2.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf2.space_before = Pt(0)
    pf2.space_after = Pt(0)
    pf2.first_line_indent = Pt(0)
    run2 = p2.add_run(caption)
    run2.font.name = '宋体'
    run2.font.size = Pt(10.5)
    run2.bold = True
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def set_cell_font(cell, text, bold=False, font_size=Pt(10.5)):
    """设置表格单元格字体"""
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = font_size
    run.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def add_code_block(doc, code_text):
    """添加代码块（等宽字体，无首行缩进）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Pt(0)
    pf.left_indent = Pt(21)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def add_table_from_data(doc, headers, rows, caption=None):
    """通用表格添加函数"""
    table = doc.add_table(rows=len(rows)+1, cols=len(headers), style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        set_cell_font(table.rows[0].cells[i], h, bold=True)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            set_cell_font(table.rows[i+1].cells[j], str(val))
    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.first_line_indent = Pt(0)
        run = p.add_run(caption)
        run.font.name = '宋体'
        run.font.size = Pt(10.5)
        run.bold = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return table


# ==================== 创建文档 ====================
doc = Document()

style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(10.5)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

# ==================== 标题 ====================
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_pf = title_p.paragraph_format
title_pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
title_pf.space_before = Pt(0)
title_pf.space_after = Pt(0)
title_run = title_p.add_run('策略首秀：生产级双均线量化交易系统')
title_run.font.name = '宋体'
title_run.font.size = Pt(16)
title_run.bold = True
title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_pf = sub_p.paragraph_format
sub_pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
sub_pf.space_before = Pt(0)
sub_pf.space_after = Pt(0)
sub_run = sub_p.add_run('TASK3 作业报告（生产级回测框架）')
sub_run.font.name = '宋体'
sub_run.font.size = Pt(14)
sub_run.bold = True
sub_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

name_p = doc.add_paragraph()
name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
name_pf = name_p.paragraph_format
name_pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
name_pf.space_before = Pt(0)
name_pf.space_after = Pt(0)
name_run = name_p.add_run('姓名：林富强')
name_run.font.name = '宋体'
name_run.font.size = Pt(10.5)
name_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

add_paragraph(doc, '', first_line_indent=False)

# ==================== 一、交易策略四要素设计 ====================
add_heading(doc, '一、交易策略四要素设计')

add_paragraph(doc,
    '一个可落地的量化交易策略，绝非简单的"金叉买、死叉卖"。真正的生产级策略必须同时回答四个问题：'
    '什么情况下买入？什么情况下卖出？每次投入多少资金？如何控制亏损？这四个问题分别对应入场规则、出场规则、仓位管理和风险控制。'
    '本报告围绕长江电力（600900.SH）近一年日线数据，构建了一套可配置、可扩展、可回测验证的生产级双均线交易系统。')

add_heading(doc, '1.1 入场规则：明确、可量化、可回测')
add_paragraph(doc,
    '好的入场规则必须具备三个特征：条件明确、可量化、可用历史数据回测。相比"感觉股票要涨了就买入"这类主观判断，'
    '本系统采用"5日均线上穿20日均线，且RSI低于70、成交量放量"作为入场条件。')
add_paragraph(doc,
    '具体规则如下：'
    '（1）均线金叉：当日MA5大于MA20，且前一日MA5小于或等于MA20；'
    '（2）RSI过滤：14日RSI小于70，避免在超买区域追买；'
    '（3）成交量过滤：当日成交量不低于20日成交量的均值，确认资金关注度。'
    '只有三个条件同时满足，才触发买入信号。')
add_paragraph(doc,
    '此外，系统还预留了趋势强度过滤（收盘价位于长均线上方）等扩展条件，便于后续根据标的特性进一步打磨。')

add_heading(doc, '1.2 出场规则：止盈、止损、信号、时间四位一体')
add_paragraph(doc,
    '出场比入场更重要，它直接决定盈亏。本系统同时实现了四种出场规则，任一条件触发即执行卖出：')
add_paragraph(doc,
    '（1）止盈出场：当浮动收益达到10%时，锁定利润；'
    '（2）止损出场：当浮动亏损达到5%时，截断亏损；'
    '（3）信号出场：当MA5下穿MA20形成死叉时，跟随趋势反转离场；'
    '（4）时间出场：持仓超过20个交易日，无论盈亏均强制离场，避免长期套牢和资金占用。')
add_paragraph(doc,
    '四种出场规则相互补充：止盈保护利润，止损控制单笔风险，信号出场跟随趋势，时间出场防止死扛。'
    '从回测结果看，长江电力基准策略4次交易中，3次由死叉信号出场，1次由时间出场触发，未触碰到止盈或止损，'
    '说明该阶段价格以窄幅震荡为主，波动未达预设阈值。')

add_heading(doc, '1.3 仓位管理：四种模式可切换')
add_paragraph(doc,
    '仓位管理决定最大亏损，而非最大收益。本系统支持四种仓位管理模式，可根据策略风险偏好和标的特性灵活切换：')
add_paragraph(doc,
    '（1）固定金额模式：每次买入固定金额（如1万元），简单直观，适合初学者；'
    '（2）固定比例模式：每次投入总资金的固定比例（如25%），风险敞口稳定；'
    '（3）凯利公式模式：根据历史胜率和盈亏比动态调整仓位，采用半凯利（Half-Kelly）以降低参数敏感性；'
    '（4）风险平价模式：根据ATR波动率调整仓位，波动大时减仓，波动小时加仓。')
add_paragraph(doc,
    '默认采用半凯利公式。在交易记录不足5笔时，使用25%默认仓位；历史交易足够后，按 f = 0.5 × (W - (1-W)/R) 计算，'
    '其中W为胜率，R为盈亏比。仓位被严格限制在5%至50%之间，避免极端仓位。')

add_heading(doc, '1.4 风险控制：单笔、总仓位、回撤、分散化四层防线')
add_paragraph(doc,
    '风险控制是策略生存的生命线。本系统设置四层防线：')
add_paragraph(doc,
    '（1）单笔止损：单次交易最大亏损不超过总资金的2%，通过仓位与止损幅度的联动控制；'
    '（2）总仓位控制：任何时刻总持仓不超过资金的80%，保留流动性；'
    '（3）最大回撤控制：当累计回撤超过20%或单日亏损超过5%时，暂停新开仓，防止情绪化和连续亏损；'
    '（4）分散化：本框架可扩展至多个标的，避免把所有资金押注单一股票。')

add_paragraph(doc, '', first_line_indent=False)

# ==================== 二、量化策略绩效评估指标 ====================
add_heading(doc, '二、量化策略绩效评估指标')

add_paragraph(doc,
    '评价交易策略必须综合收益、风险、交易质量等多维度指标。本系统计算七大核心指标，并额外补充索提诺比率、年化波动率、Calmar比率等。')

add_heading(doc, '2.1 收益类指标')
add_paragraph(doc,
    '总收益率表示回测期间组合净值相对于初始资金的总变化。'
    '年化收益率将总收益率按交易日数折算为一年期收益，便于跨期限比较。'
    '超额收益衡量策略相对买入持有的额外收益，是判断主动交易是否创造价值的关键。')

add_heading(doc, '2.2 风险类指标')
add_paragraph(doc,
    '最大回撤（MDD）定义为组合净值从历史最高点到后续最低点的最大跌幅，是衡量极端风险的核心指标。'
    '年化波动率反映策略收益的波动程度，波动越低，持有体验越稳定。')

add_heading(doc, '2.3 交易质量类指标')
add_paragraph(doc,
    '胜率指盈利交易次数占总交易次数的比例，反映策略发出正确信号的概率。'
    '盈亏比指平均盈利幅度与平均亏损幅度的比值，盈亏比大于1意味着即使胜率不高，策略仍可能通过"让利润奔跑、截断亏损"获利。')

add_heading(doc, '2.4 综合类指标')
add_paragraph(doc,
    '夏普比率衡量每承担一单位风险所获得的超额收益，无风险利率取2%。'
    '索提诺比率只考虑下行波动，更适合评估策略的亏损风险。'
    'Calmar比率是年化收益率与最大回撤绝对值的比值，将收益与极端风险直接挂钩。')

add_paragraph(doc, '', first_line_indent=False)

# ==================== 三、真实交易成本模型 ====================
add_heading(doc, '三、真实交易成本模型')

add_paragraph(doc,
    '生产级回测必须纳入真实交易成本。本系统采用的成本模型包括：'
    '滑点0.1%（买卖均影响）、佣金0.03%（最低5元，买卖均收取）、印花税0.05%（仅卖出收取）。'
    '实际买入价 = 收盘价 × (1 + 滑点)，实际卖出价 = 收盘价 × (1 - 滑点)。')
add_paragraph(doc,
    f'在长江电力MA5/MA20基准策略中，共完成{base["n_trades"]}次交易，总交易成本为{base["total_cost"]:.2f}元，'
    f'占初始资金约{base["total_cost"]/base["initial_capital"]*100:.2f}%。'
    '虽然单次成本不大，但在高换手策略中，成本对收益的侵蚀会显著放大。')

add_paragraph(doc, '', first_line_indent=False)

# ==================== 四、Python编程实现 ====================
add_heading(doc, '四、Python编程实现')

add_heading(doc, '4.1 技术指标计算')
add_paragraph(doc,
    '系统首先计算移动平均线、RSI、成交量均线、ATR和波动率。其中RSI用于过滤超买，成交量用于确认信号质量，ATR用于风险平价仓位。核心代码如下：')
add_code_block(doc,
    "df['MA5'] = df['close'].rolling(window=5, min_periods=5).mean()\n"
    "df['MA20'] = df['close'].rolling(window=20, min_periods=20).mean()\n"
    "df['RSI'] = calc_rsi(df['close'], 14)\n"
    "df['vol_ma'] = df['vol'].rolling(window=20, min_periods=20).mean()\n"
    "df['ATR'] = calc_atr(df, 14)")

add_heading(doc, '4.2 入场信号生成')
add_paragraph(doc,
    '入场信号通过多条件过滤生成，只有金叉、RSI未超买、成交量放量三个条件同时满足时才触发。代码如下：')
add_code_block(doc,
    "golden_cross = (df['ma_diff'] > 0) & (df['ma_diff'].shift(1) <= 0)\n"
    "filters = golden_cross\n"
    "filters &= (df['RSI'] < 70)\n"
    "filters &= (df['vol'] >= df['vol_ma'])\n"
    "df.loc[filters, 'signal'] = 1")

add_heading(doc, '4.3 交易信号可视化')
add_paragraph(doc,
    '将收盘价、MA5、MA20绘制在同一图中，用红色上三角标记买入信号，绿色下三角标记卖出信号。')
add_image(doc, img_path('5_20', 'strategy.png'),
          '图1  双均线策略（MA5/MA20）交易信号图', width=Inches(6.2))
add_paragraph(doc,
    '图1中，黑色实线为收盘价，红色为MA5，蓝色为MA20。由于入场加入了RSI和成交量过滤，实际买入信号数量少于单纯金叉点；'
    '出场则由止盈、止损、死叉、时间四个规则共同决定，图中绿色下三角标记了触发卖出信号的交易日。')

add_heading(doc, '4.4 模拟回测与净值走势')
add_paragraph(doc,
    '回测以10万元初始资金开始，遵循上述入场、出场、仓位和风控规则，每日计算组合价值，并与买入持有策略对比。')
add_image(doc, img_path('5_20', 'backtest.png'),
          '图2  策略回测净值曲线（MA5/MA20）', width=Inches(6.2))
add_paragraph(doc,
    f'图2展示了双均线策略与买入持有的净值对比。回测期间策略最终价值为{base["final_value"]:,.2f}元，'
    f'总收益率{base["total_return"]}%，买入持有收益率为{base["buy_hold_return"]}%，'
    f'超额收益为{base["excess_return"]}%。在整体下跌的市场中，策略通过及时离场和仓位控制，显著跑赢了买入持有。')

add_heading(doc, '4.5 最大回撤与出场原因分析')
add_image(doc, img_path('5_20', 'drawdown.png'),
          '图3  策略最大回撤（MDD）分析', width=Inches(6.2))
add_paragraph(doc,
    f'图3展示了策略每日回撤幅度。回测期间最大回撤为{base["max_drawdown"]}%，'
    '由于仓位管理和止损规则的存在，回撤被控制在极低水平。')
add_image(doc, img_path('5_20', 'exit_reasons.png'),
          '图4  出场原因分布', width=Inches(5.0))
add_paragraph(doc,
    f'图4显示，{base["n_trades"]}次交易中，{base["exit_reasons"].get("signal", 0)}次由死叉信号触发，'
    f'{base["exit_reasons"].get("take_profit", 0)}次止盈，{base["exit_reasons"].get("stop_loss", 0)}次止损，'
    f'{base["exit_reasons"].get("time_exit", 0)}次时间出场。')

add_heading(doc, '4.6 绩效指标雷达图')
add_image(doc, img_path('5_20', 'metrics_radar.png'),
          '图5  策略绩效雷达图（归一化）', width=Inches(5.0))
add_paragraph(doc,
    '图5将总收益、年化收益、夏普、胜率、盈亏比、Calmar六大指标归一化展示。'
    '由于市场整体下跌，收益类指标偏向内侧，但风险控制指标表现较好。')

add_heading(doc, '4.7 核心指标汇总')
add_table_from_data(doc,
    ['指标', '数值'],
    [
        ['初始资金', f'{base["initial_capital"]:,.0f} 元'],
        ['最终组合价值', f'{base["final_value"]:,.2f} 元'],
        ['总收益率', f'{base["total_return"]}%'],
        ['年化收益率', f'{base["annual_return"]}%'],
        ['买入持有收益率', f'{base["buy_hold_return"]}%'],
        ['超额收益', f'{base["excess_return"]}%'],
        ['最大回撤', f'{base["max_drawdown"]}%'],
        ['年化波动率', f'{base["volatility"]}%'],
        ['夏普比率', f'{base["sharpe_ratio"]}'],
        ['索提诺比率', f'{base["sortino_ratio"]}'],
        ['Calmar 比率', f'{base["calmar_ratio"]}'],
        ['胜率', f'{base["win_rate"]}%'],
        ['盈亏比', f'{base["profit_loss_ratio"]}'],
        ['交易次数', f'{base["n_trades"]} 次'],
        ['平均持仓天数', f'{base["avg_holding_days"]} 天'],
        ['总交易成本', f'{base["total_cost"]:,.2f} 元'],
    ],
    caption='表1  MA5/MA20双均线策略核心绩效指标汇总')
add_paragraph(doc,
    f'从表1可以看出，生产级策略在长江电力上的总收益为{base["total_return"]}%，超额收益为{base["excess_return"]}%，'
    f'优于买入持有。最大回撤控制在{base["max_drawdown"]}%以内，体现了仓位管理与出场规则对风险的有效抑制。'
    f'胜率为{base["win_rate"]}%，盈亏比为{base["profit_loss_ratio"]}，说明该阶段策略盈利与亏损交替出现。')

add_paragraph(doc, '', first_line_indent=False)

# ==================== 五、交易明细 ====================
add_heading(doc, '五、交易明细')
add_paragraph(doc,
    '下表列出了MA5/MA20策略在长江电力上的全部交易记录，包括买入日期、卖出日期、持仓天数、成交股数、'
    '毛盈亏、交易成本、净盈亏、净收益率和出场原因。')

trade_rows = []
for t in trade_records:
    reason_map = {'signal': '信号', 'take_profit': '止盈', 'stop_loss': '止损',
                  'time_exit': '时间', 'final': '期末'}
    trade_rows.append([
        t['entry_date'], t['exit_date'], t['holding_days'],
        t['shares'], f"{t['gross_pnl']:.2f}", f"{t['cost']:.2f}",
        f"{t['net_pnl']:.2f}", f"{t['net_return']:.2f}%",
        reason_map.get(t['exit_reason'], t['exit_reason'])
    ])
add_table_from_data(doc,
    ['买入日期', '卖出日期', '持仓天数', '股数', '毛盈亏', '成本', '净盈亏', '净收益率', '出场原因'],
    trade_rows,
    caption='表2  MA5/MA20策略交易明细')
add_paragraph(doc,
    '从表2可以看出，尽管市场整体下跌，但单笔亏损均被严格限制。'
    '例如最后一笔交易虽然亏损超过1000元，但相对于10万元本金，亏损幅度约为1%，远小于买入持有的跌幅。')

add_paragraph(doc, '', first_line_indent=False)

# ==================== 六、不同维度对比分析 ====================
add_heading(doc, '六、不同维度对比分析')

add_heading(doc, '6.1 不同仓位管理模式对比')
add_paragraph(doc,
    '为研究仓位管理对策略的影响，在相同的入场、出场和风控规则下，分别测试固定金额、固定比例、凯利公式和风险平价四种模式。')

position_rows = []
for r in position_results:
    position_rows.append([
        r['position_mode'],
        f'{r["total_return"]}', f'{r["buy_hold_return"]}', f'{r["excess_return"]}',
        f'{r["max_drawdown"]}', f'{r["sharpe_ratio"]}', f'{r["win_rate"]}',
        f'{r["profit_loss_ratio"]}', f'{r["n_trades"]}'
    ])
add_table_from_data(doc,
    ['仓位模式', '总收益(%)', '买入持有(%)', '超额收益(%)', '最大回撤(%)', '夏普', '胜率(%)', '盈亏比', '交易次数'],
    position_rows,
    caption='表3  不同仓位管理模式绩效对比')

add_image(doc, img_path('', 'comparison_positions.png'),
          '图6  不同仓位管理模式综合对比', width=Inches(6.2))
add_paragraph(doc,
    '图6显示，固定金额模式回撤最小（约0.72%），但资金利用率较低；凯利公式和固定比例模式表现接近；'
    '风险平价模式因根据ATR放大波动敞口，回撤和成本略高。 overall，四种模式均显著跑赢了买入持有。')

add_heading(doc, '6.2 不同均线周期对比')
add_paragraph(doc,
    '为研究均线周期的影响，选取五组常用组合在长江电力上回测。')

period_rows = []
for r in period_results:
    period_rows.append([
        f'MA{r["short_period"]}/MA{r["long_period"]}',
        f'{r["total_return"]}', f'{r["buy_hold_return"]}', f'{r["excess_return"]}',
        f'{r["max_drawdown"]}', f'{r["sharpe_ratio"]}', f'{r["win_rate"]}',
        f'{r["profit_loss_ratio"]}', f'{r["n_trades"]}'
    ])
add_table_from_data(doc,
    ['均线组合', '总收益(%)', '买入持有(%)', '超额收益(%)', '最大回撤(%)', '夏普', '胜率(%)', '盈亏比', '交易次数'],
    period_rows,
    caption='表4  不同均线周期绩效对比')

add_image(doc, img_path('', 'comparison_periods.png'),
          '图7  不同均线周期综合对比', width=Inches(6.2))
add_paragraph(doc,
    '从表4和图7可以看出，MA5/MA20和MA10/MA20表现较为均衡，超额收益分别为5.21%和5.35%，最大回撤控制在1.5%至1.8%之间。'
    'MA5/MA10交易次数最多（5次），胜率较低；MA5/MA15则在该阶段4次交易全部亏损。'
    '这说明参数选择对策略表现影响显著，需结合标的波动特性进行优化。')

add_heading(doc, '6.3 不同股票对比')
add_paragraph(doc,
    '为研究策略在不同股票上的适用性，选取长江电力、贵州茅台、中国平安三只股票进行对比。')

stock_rows = []
for r in stock_results:
    stock_rows.append([
        r['name'], f'{r["total_return"]}', f'{r["buy_hold_return"]}', f'{r["excess_return"]}',
        f'{r["max_drawdown"]}', f'{r["sharpe_ratio"]}', f'{r["win_rate"]}',
        f'{r["profit_loss_ratio"]}', f'{r["n_trades"]}'
    ])
add_table_from_data(doc,
    ['股票', '总收益(%)', '买入持有(%)', '超额收益(%)', '最大回撤(%)', '夏普', '胜率(%)', '盈亏比', '交易次数'],
    stock_rows,
    caption='表5  不同股票双均线策略（MA5/MA20）对比')

add_image(doc, img_path('', 'comparison_stocks.png'),
          '图8  不同股票综合对比', width=Inches(6.2))
add_paragraph(doc,
    '表5显示，贵州茅台和中国平安的超额收益分别高达14.30%和13.73%，远超长江电力的5.21%。'
    '这验证了双均线策略更适合高波动标的：高波动股票趋势更明显，止损和仓位管理更能发挥风险控制价值。'
    '但值得注意的是，贵州茅台和中国平安在回测期内触发交易次数极少（各1次），样本量有限，结论需谨慎看待。')

add_paragraph(doc, '', first_line_indent=False)

# ==================== 七、总结与心得 ====================
add_heading(doc, '七、总结与心得')

add_paragraph(doc,
    '通过本次TASK3的实践，对生产级量化策略的设计有了更深入的理解。以下从四要素角度总结心得：')

add_paragraph(doc,
    '第一，入场规则必须明确可量化。主观判断无法回测，而"金叉+RSI+成交量"这类组合条件可以客观验证，'
    '也能通过历史数据筛选出更优的过滤条件。')

add_paragraph(doc,
    '第二，出场规则比入场更重要。止盈锁定利润、止损截断亏损、信号跟随趋势、时间防止套牢，四位一体的出场机制'
    '让策略在不同行情下都有退出路径，避免出现"买得对、卖不对"的困境。')

add_paragraph(doc,
    '第三，仓位管理决定最大亏损。固定金额、固定比例、凯利公式、风险平价各有优劣：'
    '固定金额简单稳健，凯利公式数学最优但参数敏感，风险平价与波动率挂钩。'
    '生产环境中建议结合实际资金规模、标的相关性和风险预算综合选择。')

add_paragraph(doc,
    '第四，风险控制是生存底线。单笔止损、总仓位上限、最大回撤熔断、单日大跌熔断四层防线，'
    '能够有效防止极端行情下的 catastrophic loss。')

add_paragraph(doc,
    '第五，交易成本不可忽视。滑点、佣金、印花税在频繁交易中会显著侵蚀收益。'
    '纳入真实成本后，部分理论上盈利的策略可能变得不盈利，因此回测必须贴近真实交易环境。')

add_paragraph(doc,
    '第六，策略与标的特性匹配。高波动股票更适合趋势跟踪策略，低波动股票信号稀疏、收益有限。'
    '未来可进一步引入波动率过滤，只在趋势明显的市场中开仓。')

add_paragraph(doc,
    '第七，综合评估优于单一指标。总收益、最大回撤、夏普、胜率、盈亏比、Calmar等指标相互关联，'
    '必须结合解读。本次生产级框架在控制回撤的前提下，实现了稳定的超额收益，为后续更复杂策略研究奠定了坚实基础。')

# ==================== 保存文档 ====================
docx_path = os.path.join(BASE_DIR, '林富强+TASK3.docx')
doc.save(docx_path)
print(f'Word文档已保存至: {docx_path}')

# ==================== 转换为PDF ====================
try:
    from docx2pdf import convert
    pdf_path = os.path.join(BASE_DIR, '林富强+TASK3.pdf')
    convert(docx_path, pdf_path)
    print(f'PDF文档已保存至: {pdf_path}')
except Exception as e:
    print(f'PDF转换失败: {e}')
    print('请手动将docx转换为PDF，或检查Microsoft Word是否已安装。')

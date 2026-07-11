# -*- coding: utf-8 -*-
"""
生成 TASK4 海龟策略作业报告（Word + PDF）。
格式要求：宋体，五号字，1.5倍行距，0倍段间距，文字两端对齐。
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import json

# ==================== 报告路径 ====================
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DATA_DIR = os.path.join(BASE_DIR, 'data')
IMAGE_DIR = os.path.join(BASE_DIR, 'images')


def set_cell_font(cell, font_name='宋体', font_size=Pt(10.5), bold=False):
    """设置表格单元格字体"""
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = font_size
            run.bold = bold
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def add_paragraph_with_format(doc, text, font_name='宋体', font_size=Pt(10.5),
                               bold=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                               first_line_indent=True, space_before=Pt(0), space_after=Pt(0)):
    """添加带格式的段落"""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = space_before
    pf.space_after = space_after
    if first_line_indent:
        pf.first_line_indent = Pt(21)
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = font_size
    run.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return p


def add_heading_custom(doc, text, level=1):
    """添加自定义标题（宋体五号加粗）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Pt(0)
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(10.5)
    run.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p


def add_image_with_caption(doc, image_path, caption, width=Cm(14)):
    """添加图片及标题"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_picture(image_path, width=width)
    # 图片标题
    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    cap_p.paragraph_format.space_before = Pt(0)
    cap_p.paragraph_format.space_after = Pt(6)
    cap_run = cap_p.add_run(caption)
    cap_run.font.name = '宋体'
    cap_run.font.size = Pt(10)
    cap_run.bold = True
    cap_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def load_results():
    """加载回测结果 JSON"""
    results_path = os.path.join(DATA_DIR, 'backtest_results.json')
    if not os.path.exists(results_path):
        return None
    with open(results_path, 'r', encoding='utf-8') as f:
        return json.load(f)


# ==================== 创建文档 ====================
doc = Document()

# 设置默认字体
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(10.5)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 设置页边距
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

# ==================== 标题 ====================
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_pf = title_p.paragraph_format
title_pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
title_pf.space_before = Pt(0)
title_pf.space_after = Pt(0)
title_run = title_p.add_run('海龟交易策略：唐奇安通道与 ATR 动态止损')
title_run.font.name = '宋体'
title_run.font.size = Pt(16)
title_run.bold = True
title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 副标题
sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_pf = sub_p.paragraph_format
sub_pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
sub_pf.space_before = Pt(0)
sub_pf.space_after = Pt(0)
sub_run = sub_p.add_run('TASK4 作业报告')
sub_run.font.name = '宋体'
sub_run.font.size = Pt(14)
sub_run.bold = True
sub_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 姓名信息
name_p = doc.add_paragraph()
name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
name_pf = name_p.paragraph_format
name_pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
name_pf.space_before = Pt(0)
name_pf.space_after = Pt(0)
name_run = name_p.add_run('姓名：林富强')
name_run.font.name = '宋体'
name_run.font.size = Pt(10.5)
name_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

add_paragraph_with_format(doc, '', first_line_indent=False)

# ==================== 一、海龟策略核心思想与优势 ====================
add_heading_custom(doc, '一、海龟交易策略核心思想与关键优势')

add_paragraph_with_format(doc,
    '海龟交易策略（Turtle Trading System）起源于1983年美国期货交易员理查德·丹尼斯'
    '（Richard Dennis）与威廉·埃克哈特（William Eckhardt）的著名实验。丹尼斯认为'
    '交易技能可以通过系统化训练习得，而非天赋。他招募了一批毫无交易经验的学员，'
    '教授他们一套完整的趋势跟踪交易系统，这些学员后来被称为"海龟"。实验结果证明，'
    '这套规则化的策略在多种期货市场中取得了显著的长期收益。')

add_paragraph_with_format(doc,
    '海龟策略的核心思想可以概括为：跟随趋势、严格纪律、科学风控。具体而言，'
    '它是一套完整的趋势跟踪交易系统，通过识别价格突破来捕捉市场趋势，'
    '利用平均真实波幅（ATR）进行动态仓位管理和止损控制，并通过严格的纪律确保'
    '交易规则的一致执行。')

add_heading_custom(doc, '1.1 海龟策略的关键优势')

add_paragraph_with_format(doc,
    '（1）规则化交易，消除情绪干扰：海龟策略将交易决策完全规则化，'
    '从入场、加仓、止损到出场都有明确的数学规则。交易者不需要判断市场方向，'
    '只需机械执行规则，从而避免了贪婪、恐惧等情绪对交易决策的干扰。')

add_paragraph_with_format(doc,
    '（2）趋势跟踪，捕捉大行情：策略基于价格突破设计，天然具有趋势跟踪属性。'
    '当市场形成明显趋势时，策略能够及时介入并持有，直到趋势反转。这种'
    '让利润奔跑（Let Profits Run）的哲学使策略在趋势行情中能够获取显著收益。')

add_paragraph_with_format(doc,
    '（3）动态仓位管理，精细化风险控制：通过 ATR 计算单位头寸，'
    '确保每个单位的风险恒定（通常为账户资金的1%）。市场波动大时自动减少仓位，'
    '波动小时自动增加仓位，实现了风险暴露的动态平衡。')

add_paragraph_with_format(doc,
    '（4）金字塔加仓，放大盈利：在趋势确认后，策略允许最多加仓至4个单位，'
    '每上涨0.5倍ATR增加一个单位。这种金字塔式加仓在趋势延续时放大盈利，'
    '同时由于每个单位的止损独立计算，整体风险仍然可控。')

add_paragraph_with_format(doc,
    '（5）适应性止损，保护利润：动态止损价跟随入场价和ATR调整，'
    '波动大时止损距离放宽（避免正常波动误触发），波动小时止损距离收紧（更好保护利润）。'
    '这种自适应机制提高了止损的有效性。')

add_paragraph_with_format(doc,
    '（6）严格的纪律性：策略要求所有信号必须执行，触发止损必须立即离场，'
    '不留犹豫空间。这种铁一般的纪律是海龟策略能够在长期中生存并获得正收益的'
    '根本保障。')

add_paragraph_with_format(doc, '', first_line_indent=False)

# ==================== 二、核心概念解释 ====================
add_heading_custom(doc, '二、海龟策略核心概念解释')

add_heading_custom(doc, '2.1 唐奇安通道（Donchian Channel）——高低价格通道')

add_paragraph_with_format(doc,
    '唐奇安通道是海龟策略的入场和出场核心工具，由理查德·唐奇安（Richard Donchian）'
    '提出。它由两条线组成：上轨为过去N个交易日的最高价，下轨为过去N个交易日的最低价。'
    '当价格突破上轨时，视为潜在的上涨趋势启动信号；当价格跌破下轨时，视为趋势反转或结束信号。')

add_paragraph_with_format(doc,
    '海龟策略设计了两个系统：'
    '系统一（S1）使用较短的周期（突破20日最高价买入，跌破10日最低价卖出），'
    '更敏感、信号更多、但假突破也较多；'
    '系统二（S2）使用较长的周期（突破55日最高价买入，跌破20日最低价卖出），'
    '更稳健、过滤更多噪音、但信号更少。通常两个系统同时运行以分散风险。')

add_heading_custom(doc, '2.2 平均真实波幅（ATR）——动态波动度量')

add_paragraph_with_format(doc,
    '平均真实波幅（Average True Range, ATR）由J. Welles Wilder Jr.提出，'
    '是衡量市场波动性的核心指标。真实波幅（TR）取以下三者中的最大值：'
    '（1）当日最高价 - 当日最低价；'
    '（2）当日最高价 - 前日收盘价的绝对值；'
    '（3）当日最低价 - 前日收盘价的绝对值。'
    'ATR 则是真实波幅的N日简单移动平均（海龟策略通常使用20日ATR）。')

add_paragraph_with_format(doc,
    'ATR 在海龟策略中承担三个关键角色：'
    '（1）仓位计算：单位头寸 =（账户资金 × 1%）/ ATR，'
    '确保每个单位的风险恒定；'
    '（2）止损设定：止损价 = 入场价 - 2 × ATR，'
    '给予策略足够的呼吸空间，同时保护资本；'
    '（3）加仓触发：每上涨0.5 × ATR增加一个单位，'
    '在趋势延续中逐步放大盈利。')

add_heading_custom(doc, '2.3 止损条件——双重保护机制')

add_paragraph_with_format(doc,
    '海龟策略的止损系统包含两个层次：'
    '（1）固定ATR止损：每个单位建立时设定止损价 = 入场价 - 2 × ATR。'
    '如果价格跌破止损价，立即平仓离场。该止损是动态计算的，'
    '随市场波动自动调整，波动大时止损距离放宽，波动小时收紧。'
    '（2）反向突破出场：当价格跌破 exit_window 日最低价时，'
    '视为趋势结束信号，无论当前盈亏全部清仓。'
    '这一机制确保在趋势反转时及时离场，避免利润大幅回吐。')

add_paragraph_with_format(doc,
    '两个出场条件的关系：ATR止损优先保护本金，防止单笔大幅亏损；'
    '反向突破出场则优先保护利润，让盈利头寸尽可能长时间地增长。'
    '在实际运行中， whichever comes first 即触发。')

add_paragraph_with_format(doc, '', first_line_indent=False)

# ==================== 三、Python 实现与回测结果 ====================
add_heading_custom(doc, '三、Python 实现与回测结果')

add_heading_custom(doc, '3.1 数据加载与指标计算')

add_paragraph_with_format(doc,
    '本任务使用 TASK1 中已存储的长江电力（600900.SH）日线数据，'
    '包含开盘价、最高价、最低价、收盘价、成交量等字段。'
    '通过 Python pandas 加载数据后，计算唐奇安通道（20日最高价/最低价）和 ATR（20日）。')

add_paragraph_with_format(doc,
    '图1展示了长江电价的股价走势与唐奇安通道的可视化结果。'
    '上轨（红色虚线）代表过去20日最高价，下轨（蓝色虚线）代表过去20日最低价，'
    '黄色填充区域为唐奇安通道区间。当价格向上突破上轨时，生成买入信号（红色三角标记）；'
    '当价格跌破下轨时，生成卖出信号（绿色倒三角标记）。')

# 加载结果数据
data = load_results()
base = data['base'] if data else {}
periods = data['periods'] if data else []
stocks = data['stocks'] if data else []

# 图1
img1 = os.path.join(IMAGE_DIR, '20_10', 'strategy.png')
if os.path.exists(img1):
    add_image_with_caption(doc, img1, '图1 长江电力海龟策略交易信号图（唐奇安通道 + 买卖信号）')

add_paragraph_with_format(doc, '', first_line_indent=False)

add_heading_custom(doc, '3.2 回测绩效指标')

add_paragraph_with_format(doc,
    '基于长江电力近一年数据，使用海龟策略系统一（突破20日最高买入，跌破10日最低卖出）'
    '进行回测，初始资金为100,000元。回测结果如下表所示：')

# 绩效表格
if base:
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '指标'
    hdr_cells[1].text = '数值'
    set_cell_font(hdr_cells[0], bold=True)
    set_cell_font(hdr_cells[1], bold=True)

    metrics = [
        ('初始资金', f'{base.get("initial_capital", 0):,.0f} 元'),
        ('最终资金', f'{base.get("final_value", 0):,.2f} 元'),
        ('总收益率', f'{base.get("total_return", 0):.2f}%'),
        ('年化收益率', f'{base.get("annual_return", 0):.2f}%'),
        ('买入持有收益率', f'{base.get("buy_hold_return", 0):.2f}%'),
        ('超额收益', f'{base.get("excess_return", 0):.2f}%'),
        ('最大回撤（MDD）', f'{base.get("max_drawdown", 0):.2f}%'),
        ('夏普比率（Sharpe）', f'{base.get("sharpe_ratio", 0):.3f}'),
        ('索提诺比率（Sortino）', f'{base.get("sortino_ratio", 0):.3f}'),
        ('年化波动率', f'{base.get("volatility", 0):.2f}%'),
        ('Calmar比率', f'{base.get("calmar_ratio", 0):.3f}'),
        ('胜率', f'{base.get("win_rate", 0):.2f}%'),
        ('盈亏比', f'{base.get("profit_loss_ratio", 0):.3f}'),
        ('交易次数', f'{base.get("n_trades", 0)} 次'),
        ('平均持仓天数', f'{base.get("avg_holding_days", 0)} 天'),
        ('平均加仓单位数', f'{base.get("avg_units", 0):.2f} 单位'),
        ('总交易成本', f'{base.get("total_cost", 0):.2f} 元'),
    ]
    for label, value in metrics:
        row_cells = table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = value
        set_cell_font(row_cells[0])
        set_cell_font(row_cells[1])

add_paragraph_with_format(doc,
    '从绩效指标可以看出，海龟策略在长江电力的回测中表现如何。'
    '总收益率、最大回撤、夏普比率等核心指标反映了策略的风险收益特征。')

add_paragraph_with_format(doc, '', first_line_indent=False)

add_heading_custom(doc, '3.3 回测可视化分析')

add_paragraph_with_format(doc,
    '图2展示了策略净值曲线与买入持有策略的对比。'
    '红色曲线为海龟策略净值，蓝色曲线为买入持有净值。'
    '通过对比可以观察策略是否跑赢了简单的买入持有策略。')

img2 = os.path.join(IMAGE_DIR, '20_10', 'backtest.png')
if os.path.exists(img2):
    add_image_with_caption(doc, img2, '图2 海龟策略回测净值曲线（策略 vs 买入持有）')

add_paragraph_with_format(doc, '', first_line_indent=False)

add_paragraph_with_format(doc,
    '图3展示了策略的最大回撤（MDD）分析。'
    '最大回撤是衡量策略风险承受能力的关键指标，反映了策略在历史上从峰值到谷底的最大亏损幅度。'
    '较低的回撤意味着策略具有更好的风险控制能力。')

img3 = os.path.join(IMAGE_DIR, '20_10', 'drawdown.png')
if os.path.exists(img3):
    add_image_with_caption(doc, img3, '图3 海龟策略最大回撤（MDD）分析')

add_paragraph_with_format(doc, '', first_line_indent=False)

add_paragraph_with_format(doc,
    '图4为策略绩效雷达图，将总收益、年化收益、夏普比率、胜率、盈亏比、Calmar比率'
    '六个核心指标进行归一化展示，便于直观了解策略的综合表现。')

img4 = os.path.join(IMAGE_DIR, '20_10', 'metrics_radar.png')
if os.path.exists(img4):
    add_image_with_caption(doc, img4, '图4 海龟策略绩效指标雷达图（归一化）')

add_paragraph_with_format(doc, '', first_line_indent=False)

add_paragraph_with_format(doc,
    '图5为出场原因分布饼图，展示了交易以何种方式结束：止损出场、反向突破出场或期末清仓。'
    '分析出场原因有助于理解策略的行为特征和潜在的改进方向。')

img5 = os.path.join(IMAGE_DIR, '20_10', 'exit_reasons.png')
if os.path.exists(img5):
    add_image_with_caption(doc, img5, '图5 海龟策略出场原因分布')

add_paragraph_with_format(doc, '', first_line_indent=False)

# ==================== 四、参数调节与对比分析 ====================
add_heading_custom(doc, '四、参数调节与对比分析')

add_heading_custom(doc, '4.1 不同通道周期对比')

add_paragraph_with_format(doc,
    '为了探究通道周期对策略性能的影响，本任务测试了八组不同的入场/出场周期组合：'
    '（1）极短周期：入场5日/出场3日；'
    '（2）短周期对称：入场5日/出场5日；'
    '（3）短周期：入场10日/出场5日；'
    '（4）短周期对称：入场10日/出场10日；'
    '（5）中短周期：入场15日/出场7日；'
    '（6）系统一：入场20日/出场10日；'
    '（7）中长周期：入场30日/出场15日；'
    '（8）系统二：入场55日/出场20日。'
    '各组参数在相同股票（长江电力）和相同时间段上进行回测，结果如图6和表1所示。')

img6 = os.path.join(IMAGE_DIR, 'comparison_periods.png')
if os.path.exists(img6):
    add_image_with_caption(doc, img6, '图6 海龟策略不同通道周期综合对比')

if periods:
    add_paragraph_with_format(doc, '表1 不同通道周期回测绩效对比', first_line_indent=False)
    t = doc.add_table(rows=1, cols=8)
    t.style = 'Table Grid'
    hdr = t.rows[0].cells
    headers = ['周期组合', '总收益', '年化收益', '超额', '最大回撤', '夏普', '胜率', '交易次数']
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_font(hdr[i], bold=True)
    for r in periods:
        row = t.add_row().cells
        row[0].text = f"入场{r['entry_window']}/出场{r['exit_window']}"
        row[1].text = f"{r['total_return']:.2f}%"
        row[2].text = f"{r['annual_return']:.2f}%"
        row[3].text = f"{r['excess_return']:.2f}%"
        row[4].text = f"{r['max_drawdown']:.2f}%"
        row[5].text = f"{r['sharpe_ratio']:.3f}"
        row[6].text = f"{r['win_rate']:.1f}%"
        row[7].text = str(r['n_trades'])
        for cell in row:
            set_cell_font(cell)

add_paragraph_with_format(doc,
    '从周期对比可以看出，极短周期（5/3）交易最为频繁（8次），但假突破较多，'
    '胜率仅25%，总收益-7.30%。中短周期（10/5和15/7）表现相对较好，'
    '其中10/5参数胜率达到50%，超额收益为正。经典系统一（20/10）产生3次交易，'
    '总收益-6.20%。长周期（55/20）信号最少（2次），总收益-4.81%，'
    '虽然绝对亏损最小但超额收益也为负。整体来看，在下跌市场中，'
    '中短周期参数（10/5、15/7）在捕捉反弹机会和控制风险方面表现更优。')

add_paragraph_with_format(doc, '', first_line_indent=False)

add_heading_custom(doc, '4.2 不同股票对比')

add_paragraph_with_format(doc,
    '为了检验策略的普适性，本任务在三个不同风格的股票上进行了回测：'
    '长江电力（防御型公用事业）、贵州茅台（消费型蓝筹）、中国平安（金融型蓝筹）。'
    '每只股票均测试全部八组参数，共24组回测结果，如表2所示。'
    '图7为各股票最优参数的综合对比图。')

img7 = os.path.join(IMAGE_DIR, 'comparison_stocks.png')
if os.path.exists(img7):
    add_image_with_caption(doc, img7, '图7 不同股票海龟策略综合对比（各股票最优参数）')

add_paragraph_with_format(doc, '表2 三只股票 x 八组参数完整回测结果（共24组）', first_line_indent=False)
t2 = doc.add_table(rows=1, cols=8)
t2.style = 'Table Grid'
hdr2 = t2.rows[0].cells
headers2 = ['股票', '参数', '总收益', '买入持有', '超额', '最大回撤', '胜率', '交易次数']
for i, h in enumerate(headers2):
    hdr2[i].text = h
    set_cell_font(hdr2[i], bold=True)

# 从 multi_param_results.json 读取全部24组结果，按股票分组展示
multi_param_path = os.path.join(DATA_DIR, 'multi_param_results.json')
param_order = [(5,3),(5,5),(10,5),(10,10),(15,7),(20,10),(30,15),(55,20)]
if os.path.exists(multi_param_path):
    with open(multi_param_path, 'r', encoding='utf-8') as f:
        multi_results = json.load(f)
    stock_order = ['长江电力', '贵州茅台', '中国平安']
    for stock in stock_order:
        for ew, xw in param_order:
            match = [r for r in multi_results
                     if r['stock_name'] == stock
                     and r['entry_window'] == ew
                     and r['exit_window'] == xw]
            if match:
                r = match[0]
                row = t2.add_row().cells
                row[0].text = r['stock_name']
                row[1].text = f"{r['entry_window']}/{r['exit_window']}"
                row[2].text = f"{r['total_return']:+.2f}%"
                row[3].text = f"{r['buy_hold_return']:+.2f}%"
                row[4].text = f"{r['excess_return']:+.2f}%"
                row[5].text = f"{r['max_drawdown']:.2f}%"
                row[6].text = f"{r['win_rate']:.0f}%"
                row[7].text = str(r.get('n_trades', 0))
                for cell in row:
                    set_cell_font(cell)
else:
    # 回退到原始 stock_results.json（固定20/10参数）
    for r in stocks:
        row = t2.add_row().cells
        row[0].text = r['stock_name']
        row[1].text = '20/10'
        row[2].text = f"{r['total_return']:.2f}%"
        row[3].text = f"{r['buy_hold_return']:.2f}%"
        row[4].text = f"{r['excess_return']:.2f}%"
        row[5].text = f"{r['max_drawdown']:.2f}%"
        row[6].text = f"{r['win_rate']:.1f}%"
        row[7].text = str(r.get('n_trades', 0))
        for cell in row:
            set_cell_font(cell)

add_paragraph_with_format(doc, '', first_line_indent=False)

# ==================== 4.3 多股票 x 多参数综合测试 ====================
add_heading_custom(doc, '4.3 多股票 x 多参数综合测试')

add_paragraph_with_format(doc,
    '为了更全面地评估海龟策略的适应性，本任务将三只股票与八组参数进行交叉回测，'
    '共计24组回测结果。表3展示了各股票在不同参数下的超额收益表现。')

# 加载多参数结果
multi_param_path = os.path.join(DATA_DIR, 'multi_param_results.json')
if os.path.exists(multi_param_path):
    with open(multi_param_path, 'r', encoding='utf-8') as f:
        multi_results = json.load(f)

    add_paragraph_with_format(doc, '表3 多股票 x 多参数超额收益对比（%）', first_line_indent=False)
    # 创建表格：行=参数组合，列=股票
    stock_names = ['长江电力', '贵州茅台', '中国平安']
    param_labels = ['5/3', '5/5', '10/5', '10/10', '15/7', '20/10', '30/15', '55/20']

    t3 = doc.add_table(rows=1, cols=4)
    t3.style = 'Table Grid'
    hdr3 = t3.rows[0].cells
    headers3 = ['参数组合', '长江电力', '贵州茅台', '中国平安']
    for i, h in enumerate(headers3):
        hdr3[i].text = h
        set_cell_font(hdr3[i], bold=True)

    for pl in param_labels:
        ew, xw = pl.split('/')
        row = t3.add_row().cells
        row[0].text = pl
        for j, sn in enumerate(stock_names):
            match = [r for r in multi_results
                     if r['stock_name'] == sn
                     and r['entry_window'] == int(ew)
                     and r['exit_window'] == int(xw)]
            if match:
                excess = match[0]['excess_return']
                row[j + 1].text = f"{excess:+.2f}%"
            else:
                row[j + 1].text = '-'
        for cell in row:
            set_cell_font(cell)

add_paragraph_with_format(doc,
    '从综合测试结果可以得出以下关键发现：'
    '（1）中国平安是海龟策略表现最佳的标的，在全部8组参数下均获得正超额收益，'
    '其中20/10参数下策略实现正收益+2.43%，超额收益高达+18.40%，胜率50%，盈亏比4.69。'
    '这是因为中国平安在回测期内波动较大，出现了多次可捕捉的趋势行情。'
    '（2）贵州茅台在长周期参数下表现更好，55/20参数下超额收益达+12.38%，'
    '说明对于波动较大但趋势较慢的标的，长周期通道能更好过滤噪音。'
    '（3）长江电力在10/5参数下表现最优，超额收益+8.48%，胜率50%，'
    '短周期通道更适合波动较小的公用事业股。'
    '（4）总体而言，中周期参数（10/5、15/7、20/10）在多数股票上表现均衡，'
    '是海龟策略的推荐参数区间。')

img_heatmap = os.path.join(IMAGE_DIR, 'heatmap_multi_param.png')
if os.path.exists(img_heatmap):
    add_image_with_caption(doc, img_heatmap,
                           '图8 多股票 x 多参数超额收益热力图'
                           '（红色=跑赢买入持有，绿色=跑输买入持有）')
add_heading_custom(doc, '五、总结与使用心得')

add_paragraph_with_format(doc,
    '通过本次 TASK4 的编程实践，我对海龟交易策略有了更深入的理解。'
    '海龟策略并非一个"完美"的策略，但它提供了一套完整、可执行、可验证的交易框架，'
    '其精髓在于纪律性和风险控制的系统化。')

add_paragraph_with_format(doc,
    '从回测结果来看，海龟策略在趋势性较强的市场中表现优异。'
    '本次回测期内A股整体处于下跌趋势，但海龟策略通过通道突破信号和动态止损，'
    '在多数参数组合下均跑赢了买入持有策略。'
    '其中中国平安在20/10参数下实现正收益+2.43%，而同期买入持有亏损-15.96%，'
    '超额收益达+18.40%，充分体现了趋势跟踪策略在波动市场中的优势。'
    '但在震荡市中，由于频繁假突破，策略可能面临多次小额亏损，'
    '此时需要通过更严格的过滤条件或结合其他指标进行优化。')

add_paragraph_with_format(doc,
    '参数调节方面，通道周期是影响策略表现的关键因素。'
    '本次24组交叉回测显示，不同股票的最优参数并不相同：'
    '长江电力最优为10/5（超额+8.48%），贵州茅台最优为55/20（超额+12.38%），'
    '中国平安最优为20/10（超额+18.40%）。'
    '较短的周期灵敏度更高，适合波动较小的标的；较长的周期稳健性更好，适合波动较大的标的。'
    '在实际应用中，可以结合多个系统（如同时运行20/10和55/20）'
    '来分散单一参数的风险。')

add_paragraph_with_format(doc,
    'ATR 是海龟策略最精妙的部分。它不仅是仓位管理的工具，'
    '更是连接市场波动与风险控制的桥梁。'
    '通过 ATR 动态调整仓位和止损，策略实现了风险暴露的自适应，'
    '这是许多传统固定止损策略所不具备的优势。')

add_paragraph_with_format(doc,
    '总结而言，海龟策略的适应场景主要包括：'
    '（1）具有明显趋势特征的市场；'
    '（2）波动率适中且稳定的交易品种；'
    '（3）投资者能够严格执行规则的交易环境。'
    '使用心得是：策略本身只是工具，纪律执行才是生命线。'
    '无论策略多么优秀，如果不能严格执行止损和出场规则，'
    '最终都会在市场波动中遭受重创。')

# ==================== 保存文档 ====================
desktop_path = r'C:\Users\LENOVO\Desktop\quant-ai\TASK4'
docx_path = os.path.join(desktop_path, '林富强+TASK4.docx')
doc.save(docx_path)
print(f"Word文档已保存至: {docx_path}")

# ==================== 转换为PDF ====================
try:
    from docx2pdf import convert
    pdf_path = os.path.join(desktop_path, '林富强+TASK4.pdf')
    convert(docx_path, pdf_path)
    print(f"PDF文档已保存至: {pdf_path}")
except Exception as e:
    print(f"PDF转换失败: {e}")
    print("请手动将docx转换为PDF，或检查Microsoft Word是否已安装。")

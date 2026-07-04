# -*- coding: utf-8 -*-
"""
TASK2 - 生成最终作业文档
格式要求：宋体，五号字，1.5倍行距，0倍段间距，文字两端对齐
包含：数据诊断、指标说明、指标可视化、KDJ扩展指标
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import json

OUTPUT_DIR = r'C:\Users\LENOVO\Desktop\quant-ai\TASK2'

# 加载诊断摘要
with open(os.path.join(OUTPUT_DIR, 'diagnosis_summary.json'), 'r', encoding='utf-8') as f:
    summary = json.load(f)


def add_paragraph(doc, text, bold=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  first_line_indent=True, font_size=Pt(10.5)):
    """添加带格式的段落（宋体五号，1.5倍行距，0段间距，两端对齐）"""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    if first_line_indent:
        pf.first_line_indent = Pt(21)
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = font_size
    run.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p


def add_heading(doc, text, font_size=Pt(10.5)):
    """添加标题（宋体五号加粗，左对齐）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Pt(0)
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = font_size
    run.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p


def add_image(doc, img_path, caption, width=Inches(6.0)):
    """添加图片和图注"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Pt(0)
    run = p.add_run()
    run.add_picture(img_path, width=width)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf2 = p2.paragraph_format
    pf2.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf2.space_before = Pt(0)
    pf2.space_after = Pt(0)
    pf2.first_line_indent = Pt(0)
    run2 = p2.add_run(caption)
    run2.font.name = '宋体'
    run2.font.size = Pt(10.5)
    run2.bold = True
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def set_cell_font(cell, text, bold=False, font_size=Pt(10.5)):
    """设置表格单元格字体"""
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = font_size
    run.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


# ==================== 创建文档 ====================
doc = Document()

# 设置默认字体
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(10.5)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 页边距
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

# ==================== 标题 ====================
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_pf = title_p.paragraph_format
title_pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
title_pf.space_before = Pt(0)
title_pf.space_after = Pt(0)
title_run = title_p.add_run('数据炼金术：数据诊断与构造交易指标')
title_run.font.name = '宋体'
title_run.font.size = Pt(16)
title_run.bold = True
title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_pf = sub_p.paragraph_format
sub_pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
sub_pf.space_before = Pt(0)
sub_pf.space_after = Pt(0)
sub_run = sub_p.add_run('TASK2 作业报告')
sub_run.font.name = '宋体'
sub_run.font.size = Pt(14)
sub_run.bold = True
sub_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

name_p = doc.add_paragraph()
name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
name_pf = name_p.paragraph_format
name_pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
name_pf.space_before = Pt(0)
name_pf.space_after = Pt(0)
name_run = name_p.add_run('姓名：林富强')
name_run.font.name = '宋体'
name_run.font.size = Pt(10.5)
name_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

add_paragraph(doc, '', first_line_indent=False)

# ==================== 任务一：数据诊断分析 ====================
add_heading(doc, '一、数据诊断分析')

add_paragraph(doc,
    '本任务对TASK1中获取的长江电力（600900.SH）日线行情数据进行基础诊断分析，'
    '包括缺失值检查和描述性统计量计算。数据时间范围为2025年7月2日至2026年7月2日，'
    '共243个交易日，包含开盘价、最高价、最低价、收盘价、前收盘价、涨跌额、涨跌幅、'
    '成交量、成交额等11个字段。使用Python（pandas库）进行数据处理，'
    '脚本文件为data_diagnosis.py。')

add_heading(doc, '1.1 缺失值检查')

add_paragraph(doc,
    '对全部11个字段逐一进行缺失值检测，检查结果显示所有字段的缺失值数量均为0，'
    '缺失比例均为0%。这表明数据从Tushare Pro接口获取后完整保存，'
    '在传输和存储过程中未发生数据丢失，数据完整性良好，无需进行缺失值填补处理。')

# 缺失值检查表格
table = doc.add_table(rows=12, cols=3, style='Table Grid')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['字段', '缺失数量', '缺失比例(%)']
for i, h in enumerate(headers):
    set_cell_font(table.rows[0].cells[i], h, bold=True)
fields = ['ts_code', 'trade_date', 'open', 'high', 'low', 'close',
          'pre_close', 'change', 'pct_chg', 'vol', 'amount']
for i, f in enumerate(fields):
    set_cell_font(table.rows[i+1].cells[0], f)
    set_cell_font(table.rows[i+1].cells[1], '0')
    set_cell_font(table.rows[i+1].cells[2], '0.0')

add_paragraph(doc, '', first_line_indent=False)

add_heading(doc, '1.2 描述性统计量')

add_paragraph(doc,
    '对数值型字段计算描述性统计量，包括计数、均值、标准差、最小值、下四分位数、'
    '中位数、上四分位数、最大值、方差、偏度和峰度。以下列出核心字段的统计结果：')

# 描述性统计表格
stats_table = doc.add_table(rows=6, cols=8, style='Table Grid')
stats_table.alignment = WD_TABLE_ALIGNMENT.CENTER
stats_headers = ['字段', '均值', '标准差', '最小值', '中位数', '最大值', '偏度', '峰度']
for i, h in enumerate(stats_headers):
    set_cell_font(stats_table.rows[0].cells[i], h, bold=True)

stats_data = [
    ['open', '27.60', '0.95', '25.65', '27.60', '30.61', '0.88', '1.28'],
    ['high', '27.75', '0.94', '25.94', '27.70', '30.79', '0.92', '1.37'],
    ['low', '27.43', '0.95', '25.38', '27.39', '30.40', '0.84', '1.14'],
    ['close', '27.59', '0.95', '25.65', '27.61', '30.61', '0.82', '1.19'],
    ['pct_chg', '-0.033', '0.758', '-3.00', '-0.073', '2.05', '0.05', '1.30'],
]
for i, row_data in enumerate(stats_data):
    for j, val in enumerate(row_data):
        set_cell_font(stats_table.rows[i+1].cells[j], val)

add_paragraph(doc, '', first_line_indent=False)

add_paragraph(doc,
    '从描述性统计量可以看出：收盘价均值为27.59元，标准差为0.95元，'
    '价格在25.65元至30.61元之间波动，波动幅度相对适中。日涨跌幅均值为-0.033%，'
    '标准差为0.758%，偏度为0.05接近对称分布，峰度为1.30低于正态分布的3，'
    '说明日涨跌幅分布比正态分布更平坦，尾部风险较低。'
    '成交量偏度为1.69，呈明显右偏分布，存在少数放量交易日。')

add_image(doc, os.path.join(OUTPUT_DIR, 'data_diagnosis.png'),
          '图1-4  数据诊断可视化（缺失值检查、收盘价箱线图、日涨跌幅分布、成交量分布）')

add_paragraph(doc,
    '图1显示数据无缺失值，完整性良好。图2的箱线图显示收盘价主要集中在26.9至28.1元区间，'
    '存在少数高价异常点。图3的日涨跌幅直方图近似正态分布但略有右偏，大部分交易日涨跌幅'
    '在正负1%以内。图4的成交量分布呈右偏特征，大部分交易日成交量在100万手以内，'
    '个别交易日出现明显放量。')

add_paragraph(doc, '', first_line_indent=False)

# ==================== 任务二：技术指标说明 ====================
add_heading(doc, '二、技术指标说明（RSI、MACD、布林带）')

add_paragraph(doc,
    '技术指标是通过对价格、成交量等市场数据进行数学运算得到的衍生指标，'
    '用于辅助判断市场趋势和买卖时机。本部分介绍RSI、MACD和布林带三个'
    '常用技术指标的计算方法和作用。')

add_heading(doc, '2.1 RSI（相对强弱指数，Relative Strength Index）')

add_paragraph(doc,
    'RSI由技术分析大师Welles Wilder于1978年提出，是一种衡量价格动量强弱的'
    '振荡指标。其核心思想是通过比较一段时期内价格上涨幅度和下跌幅度的比值，'
    '判断市场买卖力量的强弱对比。')

add_paragraph(doc, '计算方法：', first_line_indent=True)
add_paragraph(doc,
    '第一步，计算每日价格变动值（当日收盘价减前日收盘价）。将正变动记为涨幅，'
    '负变动取绝对值记为跌幅。第二步，计算N日内（通常N=14）的平均涨幅和平均跌幅，'
    '采用Wilder平滑法。第三步，计算相对强度RS = 平均涨幅 / 平均跌幅。'
    '第四步，RSI = 100 - 100 / (1 + RS)。RSI取值范围为0至100。')

add_paragraph(doc, '作用：', first_line_indent=True)
add_paragraph(doc,
    'RSI通常以70为超买线、30为超卖线。当RSI超过70时，表明股票近期涨幅较大，'
    '可能存在超买现象，价格有回调风险；当RSI低于30时，表明股票近期跌幅较大，'
    '可能存在超卖现象，价格有反弹机会。此外，RSI还可用于识别背离信号：'
    '当价格创新高而RSI未能创新高时形成顶背离，预示上涨动力减弱；'
    '当价格创新低而RSI未能创新低时形成底背离，预示下跌动力减弱。'
    'RSI在震荡市中表现较好，在强趋势市中可能出现指标钝化。')

add_heading(doc, '2.2 MACD（指数平滑异同移动平均线）')

add_paragraph(doc,
    'MACD由Gerald Appel于1970年代提出，是利用短期和长期指数移动平均线（EMA）'
    '之间的聚合与分离关系来判断市场趋势和动量的技术指标。MACD通过快慢均线之差'
    '捕捉趋势的方向和强度变化。')

add_paragraph(doc, '计算方法：', first_line_indent=True)
add_paragraph(doc,
    '第一步，计算12日EMA（快线）和26日EMA（慢线）。EMA的计算公式为：'
    'EMA今日 = EMA昨日 + α × (收盘价 - EMA昨日)，其中α = 2/(N+1)。'
    '第二步，计算DIF（差离值）= EMA(12) - EMA(26)。'
    '第三步，计算DEA（信号线）= DIF的9日EMA。'
    '第四步，计算MACD柱状图 = (DIF - DEA) × 2。')

add_paragraph(doc, '作用：', first_line_indent=True)
add_paragraph(doc,
    'MACD主要通过DIF与DEA的交叉关系和柱状图变化来判断趋势。'
    '当DIF上穿DEA时形成"金叉"，为买入信号，表示短期趋势转强；'
    '当DIF下穿DEA时形成"死叉"，为卖出信号，表示短期趋势转弱。'
    'MACD柱状图由负转正表示多头力量增强，由正转负表示空头力量增强。'
    '此外，DIF和DEA与零轴的相对位置也具有参考意义：在零轴上方为多头市场，'
    '在零轴下方为空头市场。MACD还可用于识别顶背离和底背离，'
    '当价格创新高而MACD未能创新高时，预示趋势可能反转。')

add_heading(doc, '2.3 布林带（Bollinger Bands）')

add_paragraph(doc,
    '布林带由John Bollinger于1980年代提出，是一种基于统计原理的波动率指标。'
    '它通过移动平均线和标准差构建价格波动的上下轨道，直观展示价格的波动范围和'
    '相对位置，是判断超买超卖和波动率变化的重要工具。')

add_paragraph(doc, '计算方法：', first_line_indent=True)
add_paragraph(doc,
    '第一步，计算20日简单移动平均线（SMA）作为布林带中轨：'
    '中轨 = SMA(close, 20)。第二步，计算20日价格标准差σ。'
    '第三步，计算上轨 = 中轨 + 2 × σ，下轨 = 中轨 - 2 × σ。'
    '上轨和下轨之间的宽度反映了价格的波动率水平。')

add_paragraph(doc, '作用：', first_line_indent=True)
add_paragraph(doc,
    '布林带的主要作用包括：第一，判断超买超卖。当价格触及或突破上轨时，'
    '可能存在超买；当价格触及或跌破下轨时，可能存在超卖。第二，衡量波动率变化。'
    '布林带收窄（带宽缩小）表示波动率降低，市场可能酝酿趋势突破；'
    '布林带开口扩大（带宽增加）表示波动率上升，趋势可能加速。'
    '第三，判断趋势方向。价格在中轨上方运行视为多头趋势，在中轨下方运行为空头趋势。'
    '第四，识别价格目标。布林带上下轨可作为短期的价格目标参考。'
    '布林带在趋势市场和震荡市场均有应用价值，但需结合其他指标综合判断。')

add_paragraph(doc, '', first_line_indent=False)

# ==================== 任务三：指标计算与可视化 ====================
add_heading(doc, '三、技术指标计算与可视化')

add_paragraph(doc,
    '本部分使用Python编程实现技术指标的计算和可视化。加载TASK1中存储的'
    '长江电力（600900.SH）日线数据，分别计算RSI(14)、MACD(12,26,9)和'
    '布林带(20,2)三个技术指标，并绘制对应的可视化图形。'
    '脚本文件为calculate_indicators.py，HTML可视化报告为report.html。')

add_heading(doc, '3.1 RSI 指标计算与可视化')

add_paragraph(doc,
    '使用14日周期计算RSI指标。计算结果显示，在观察期内RSI多次触及70以上的'
    '超买区域和30附近的超卖区域，与股价的高点和低点有较好的对应关系。'
    f'最新RSI值为{summary["日涨跌幅_均值"]:.2f}附近的44.76，处于中性区间。')

add_image(doc, os.path.join(OUTPUT_DIR, 'chart_rsi.png'),
          '图5  RSI(14) 相对强弱指数可视化', width=Inches(6.2))

add_paragraph(doc,
    '图5上半部分为收盘价走势，下半部分为RSI指标曲线。图中红色虚线为超买线（70），'
    '绿色虚线为超卖线（30），灰色虚线为中轴线（50）。可以看到，当RSI进入超买区域后，'
    '股价往往出现短期回调；当RSI进入超卖区域后，股价通常出现反弹。'
    'RSI与股价之间的背离现象也多次出现，对趋势反转具有一定的预示作用。')

add_heading(doc, '3.2 MACD 指标计算与可视化')

add_paragraph(doc,
    '使用12日和26日EMA计算DIF，9日EMA计算DEA。计算结果显示，DIF与DEA在观察期内'
    '多次交叉，每次金叉后股价通常出现上涨行情，死叉后股价往往回调。')

add_image(doc, os.path.join(OUTPUT_DIR, 'chart_macd.png'),
          '图6  MACD 指标可视化', width=Inches(6.2))

add_paragraph(doc,
    '图6上半部分为收盘价走势，下半部分为MACD指标。图中橙色线为DIF（快线），'
    '蓝色线为DEA（慢线），红绿柱状图为MACD柱。红色柱表示多头力量占优，'
    '绿色柱表示空头力量占优。可以看到，MACD柱状图由负转正的时刻往往对应'
    '股价的阶段性低点，由正转负的时刻对应阶段性高点，对买卖时机的判断有参考价值。')

add_heading(doc, '3.3 布林带指标计算与可视化')

add_paragraph(doc,
    '使用20日SMA和2倍标准差计算布林带。计算结果显示，收盘价大部分时间在'
    '布林带上下轨之间运行，触及上轨时有回调压力，触及下轨时有支撑反弹。')

add_image(doc, os.path.join(OUTPUT_DIR, 'chart_boll.png'),
          '图7  布林带(Bollinger Bands)指标可视化', width=Inches(6.2))

add_paragraph(doc,
    '图7中黑色实线为收盘价，红色线为上轨，橙色虚线为中轨（20日均线），'
    '绿色线为下轨，蓝色半透明区域为布林带通道。可以看到，当布林带收窄时，'
    '市场处于相对平静状态，随后往往出现方向性突破；当布林带开口扩大时，'
    '趋势加速运行。价格在中轨上方时多头占优，在中轨下方时空头占优。')

add_paragraph(doc, '', first_line_indent=False)

# ==================== 任务四：KDJ 扩展指标 ====================
add_heading(doc, '四、扩展指标——KDJ 随机指标')

add_paragraph(doc,
    '在RSI、MACD和布林带之外，技术分析中还有许多其他典型指标，如KDJ随机指标、'
    'OBV能量潮指标、ATR平均真实波幅、CCI顺势指标、WR威廉指标等。'
    '本任务选取KDJ随机指标进行介绍、计算和可视化。')

add_heading(doc, '4.1 KDJ 指标介绍')

add_paragraph(doc,
    'KDJ指标由George C. Lane创立，是技术分析中最常用的短线超买超卖指标之一。'
    '它通过计算一定周期内（通常为9日）最高价、最低价与收盘价之间的关系，'
    '反映当前价格在近期价格区间中的相对位置，从而判断市场的超买超卖状态和'
    '短期转折信号。KDJ指标由K线、D线和J线三条曲线组成。')

add_paragraph(doc, '计算方法：', first_line_indent=True)
add_paragraph(doc,
    '第一步，计算RSV（未成熟随机值）：RSV = (收盘价 - 9日内最低价) / '
    '(9日内最高价 - 9日内最低价) × 100。第二步，计算K值：'
    'K = 2/3 × 前日K值 + 1/3 × 当日RSV（初始K值取50）。'
    '第三步，计算D值：D = 2/3 × 前日D值 + 1/3 × 当日K值（初始D值取50）。'
    '第四步，计算J值：J = 3 × K - 2 × D。')

add_paragraph(doc, '作用：', first_line_indent=True)
add_paragraph(doc,
    'KDJ取值范围理论上为0至100（J线可超出此范围）。通常以80为超买线、20为超卖线。'
    'K线上穿D线为"金叉"，是买入信号；K线下穿D线为"死叉"，是卖出信号。'
    'J线反应最为灵敏，当J值超过100时表示市场极度超买，当J值低于0时表示市场极度超卖，'
    '这些极端位置通常是短期反转的高概率信号。KDJ指标适合短线交易，在震荡市中效果较好，'
    '但在强趋势市中可能出现指标钝化现象，即KDJ长时间停留在超买或超卖区域。')

add_heading(doc, '4.2 KDJ 指标计算与可视化')

add_paragraph(doc,
    '使用9日周期计算KDJ指标。计算结果显示，K线与D线在观察期内多次交叉，'
    'J线波动最为剧烈，多次触及100以上和0以下的极端区域。')

add_image(doc, os.path.join(OUTPUT_DIR, 'chart_kdj.png'),
          '图8  KDJ 随机指标可视化', width=Inches(6.2))

add_paragraph(doc,
    '图8上半部分为收盘价走势，下半部分为KDJ指标。图中橙色线为K线，蓝色线为D线，'
    '红色线为J线。红色虚线为超买线（80），绿色虚线为超卖线（20）。'
    '可以看到，KDJ金叉往往对应股价短期低点，死叉对应短期高点。'
    'J线触及极端区域后股价出现反转的概率较高。但需要注意的是，'
    '在强趋势行情中KDJ可能出现钝化，此时应结合MACD等趋势指标综合判断。')

add_paragraph(doc, '', first_line_indent=False)

# ==================== 指标汇总 ====================
add_heading(doc, '五、技术指标最新值汇总')

add_paragraph(doc,
    '以下为长江电力（600900.SH）截至2026年7月2日的各技术指标最新值及信号判断：')

# 汇总表格
final_table = doc.add_table(rows=11, cols=4, style='Table Grid')
final_table.alignment = WD_TABLE_ALIGNMENT.CENTER
final_headers = ['指标', '关键参数', '最新值', '信号判断']
for i, h in enumerate(final_headers):
    set_cell_font(final_table.rows[0].cells[i], h, bold=True)

final_data = [
    ['RSI(14)', '超买>70, 超卖<30', '44.76', '中性'],
    ['DIF', 'MACD快线', '-0.1677', '空头'],
    ['DEA', 'MACD慢线', '-0.1131', '空头'],
    ['MACD柱', '(DIF-DEA)×2', '-0.1091', '绿柱(空)'],
    ['布林带上轨', 'SMA20+2σ', '28.26', '-'],
    ['布林带中轨', 'SMA20', '27.13', '-'],
    ['布林带下轨', 'SMA20-2σ', '26.04', '-'],
    ['K', 'KDJ K线', '46.61', '中性'],
    ['D', 'KDJ D线', '33.53', '-'],
    ['J', '3K-2D', '72.78', '正常'],
]
for i, row_data in enumerate(final_data):
    for j, val in enumerate(row_data):
        set_cell_font(final_table.rows[i+1].cells[j], val)

add_paragraph(doc, '', first_line_indent=False)

add_paragraph(doc,
    '从汇总表可以看出，截至数据期末，RSI处于中性区间，MACD的DIF和DEA均在零轴下方，'
    '显示空头趋势尚未完全扭转。布林带方面，收盘价位于中轨附近，波动率处于正常水平。'
    'KDJ的K值和D值处于中性区域，J值在正常范围内。综合各指标判断，'
    '当前长江电力股价处于震荡整理阶段，短期方向尚不明朗。')

add_paragraph(doc,
    '在实际交易中，单一指标的信号可能存在误导，建议将多个技术指标组合使用，'
    '相互验证以提高信号可靠性。例如，可以将MACD的趋势判断、RSI和KDJ的超买超卖判断、'
    '布林带的波动率边界判断结合起来，构建更完善的交易策略。同时，'
    '技术指标分析应与基本面分析相结合，在全面评估公司价值和行业前景的基础上做出投资决策。')

# ==================== 保存文档 ====================
docx_path = os.path.join(OUTPUT_DIR, '林富强+TASK2.docx')
doc.save(docx_path)
print(f'Word文档已保存至: {docx_path}')

# ==================== 转换为PDF ====================
try:
    from docx2pdf import convert
    pdf_path = os.path.join(OUTPUT_DIR, '林富强+TASK2.pdf')
    convert(docx_path, pdf_path)
    print(f'PDF文档已保存至: {pdf_path}')
except Exception as e:
    print(f'PDF转换失败: {e}')
    print('请手动将docx转换为PDF，或检查Microsoft Word是否已安装。')

# -*- coding: utf-8 -*-
"""
生成 TASK5 机器学习分类作业报告（Word + PDF）。
格式要求：宋体，五号字，1.5倍行距，0倍段间距，文字两端对齐。
五种模型：线性回归、逻辑回归、决策树、随机森林、KNN
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import json

# ==================== 路径配置 ====================
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DATA_DIR = os.path.join(BASE_DIR, 'data')
IMAGE_DIR = os.path.join(BASE_DIR, 'images')


def set_cell_font(cell, font_name='宋体', font_size=Pt(10.5), bold=False):
    """设置表格单元格字体"""
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = font_size
            run.bold = bold
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def add_paragraph_with_format(doc, text, font_name='宋体', font_size=Pt(10.5),
                               bold=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                               first_line_indent=True, space_before=Pt(0), space_after=Pt(0)):
    """添加带格式的段落"""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = space_before
    pf.space_after = space_after
    if first_line_indent:
        pf.first_line_indent = Pt(21)
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = font_size
    run.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return p


def add_heading_custom(doc, text, level=1):
    """添加自定义标题（宋体五号加粗）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Pt(0)
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(10.5)
    run.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p


def add_image_with_caption(doc, image_path, caption, width=Cm(14)):
    """添加图片及标题"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_picture(image_path, width=width)
    # 图片标题
    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    cap_p.paragraph_format.space_before = Pt(0)
    cap_p.paragraph_format.space_after = Pt(6)
    cap_run = cap_p.add_run(caption)
    cap_run.font.name = '宋体'
    cap_run.font.size = Pt(10)
    cap_run.bold = True
    cap_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def load_results():
    """加载机器学习评估结果 JSON"""
    results_path = os.path.join(DATA_DIR, 'ml_results.json')
    if not os.path.exists(results_path):
        return None
    with open(results_path, 'r', encoding='utf-8') as f:
        return json.load(f)


def add_results_table(doc, results_list, table_caption):
    """添加模型评估结果表格"""
    add_paragraph_with_format(doc, table_caption, first_line_indent=False)
    t = doc.add_table(rows=1, cols=7)
    t.style = 'Table Grid'
    hdr = t.rows[0].cells
    headers = ['模型', 'Accuracy', 'Precision', 'Recall', 'F1', 'AUC']
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_font(hdr[i], bold=True)
    for r in results_list:
        row = t.add_row().cells
        row[0].text = r['模型']
        row[1].text = f"{r['Accuracy']:.4f}"
        row[2].text = f"{r['Precision']:.4f}"
        row[3].text = f"{r['Recall']:.4f}"
        row[4].text = f"{r['F1']:.4f}"
        row[5].text = f"{r['AUC']:.4f}"
        for cell in row:
            set_cell_font(cell)
    add_paragraph_with_format(doc, '', first_line_indent=False)


# ==================== 加载结果数据 ====================
results = load_results()
cancer_results = results.get('breast_cancer', []) if results else []
stock_results = results.get('stock', []) if results else []

# ==================== 创建文档 ====================
doc = Document()

# 设置默认字体
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(10.5)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 设置页边距
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

# ==================== 标题 ====================
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_pf = title_p.paragraph_format
title_pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
title_pf.space_before = Pt(0)
title_pf.space_after = Pt(0)
title_run = title_p.add_run('AI交易引擎：机器学习算法与场景应用')
title_run.font.name = '宋体'
title_run.font.size = Pt(16)
title_run.bold = True
title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 副标题
sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_pf = sub_p.paragraph_format
sub_pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
sub_pf.space_before = Pt(0)
sub_pf.space_after = Pt(0)
sub_run = sub_p.add_run('TASK5 作业报告')
sub_run.font.name = '宋体'
sub_run.font.size = Pt(14)
sub_run.bold = True
sub_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 姓名信息
name_p = doc.add_paragraph()
name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
name_pf = name_p.paragraph_format
name_pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
name_pf.space_before = Pt(0)
name_pf.space_after = Pt(0)
name_run = name_p.add_run('姓名：林富强')
name_run.font.name = '宋体'
name_run.font.size = Pt(10.5)
name_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

add_paragraph_with_format(doc, '', first_line_indent=False)

# ==================== 摘要 ====================
add_heading_custom(doc, '摘要')

add_paragraph_with_format(doc,
    '本报告围绕机器学习算法在量化交易中的应用展开，包含三个部分：'
    '第一部分系统阐述线性回归、逻辑回归、决策树、随机森林和KNN五种经典机器学习算法的原理、优缺点及金融应用场景；'
    '第二部分解释混淆矩阵、AUC、ROC曲线等模型评价指标的概念与计算方法；'
    '第三部分通过Python编程实践，分别在scikit-learn乳腺癌数据集和长江电力（600900.SH）'
    '股票日线数据上进行二分类实验，构建并训练五种模型，计算AUC指标并绘制ROC曲线。'
    '实验结果表明，机器学习模型在结构化医学数据上表现优异（AUC最高达0.995），'
    '而在金融时间序列数据上的预测效果则受到市场有效性的挑战，'
    '这一对比结果对理解机器学习在量化交易中的适用边界具有重要参考意义。')

add_paragraph_with_format(doc, '', first_line_indent=False)

# ==================== 一、引言 ====================
add_heading_custom(doc, '一、引言')

add_heading_custom(doc, '1.1 任务背景')

add_paragraph_with_format(doc,
    '机器学习算法在量化交易领域的应用日益广泛，它能够帮助交易者从海量数据中挖掘有价值的信息，'
    '构建更精准的预测模型，从而优化交易策略。在传统的量化分析中，交易者主要依赖统计指标和技术分析'
    '进行决策，而机器学习通过自动学习数据中的复杂模式，为交易决策提供了新的方法论。')

add_paragraph_with_format(doc,
    '分类问题是机器学习中最基础也最常见的问题类型之一。在金融领域，股票涨跌方向预测、'
    '信用违约风险评估、客户流失预测等问题都可以转化为分类问题。'
    '本任务聚焦于基础的机器学习算法，通过理论与实践相结合的方式，'
    '深入理解分类算法的原理、评价指标以及在实际数据上的应用效果。')

add_heading_custom(doc, '1.2 学习目标')

add_paragraph_with_format(doc,
    '（1）理解线性回归、逻辑回归、决策树、随机森林、KNN等机器学习算法的核心思想和适用场景；'
    '（2）掌握混淆矩阵、AUC、ROC曲线等模型评价指标的计算方法和解读能力；'
    '（3）通过Python编程实践，完成从数据加载、训练集划分、模型训练到模型评估的完整流程；'
    '（4）对比分析不同模型在不同数据集上的表现差异，理解机器学习在金融数据中的应用挑战。')

add_paragraph_with_format(doc, '', first_line_indent=False)

# ==================== 二、机器学习算法 ====================
add_heading_custom(doc, '二、机器学习算法')

add_heading_custom(doc, '2.1 算法总览')

add_paragraph_with_format(doc,
    '机器学习算法按照学习任务的不同，主要分为回归、分类和聚类三大类。'
    '在量化交易领域，回归算法用于预测连续型变量（如收益率、价格），'
    '分类算法用于预测离散型标签（如涨跌方向、违约与否），'
    '聚类算法用于发现数据中的潜在分组结构。'
    '本任务重点研究五种常见的机器学习算法，下表列出了各算法及其特点。')

# 算法总览表
add_paragraph_with_format(doc, '表1 五种常见机器学习算法对比', first_line_indent=False)
t_overview = doc.add_table(rows=1, cols=4)
t_overview.style = 'Table Grid'
hdr_ov = t_overview.rows[0].cells
headers_ov = ['算法', '类型', '优势', '劣势']
for i, h in enumerate(headers_ov):
    hdr_ov[i].text = h
    set_cell_font(hdr_ov[i], bold=True)

overview_data = [
    ('线性回归', '回归', '快速、可解释', '只能捕捉线性关系'),
    ('逻辑回归', '分类', '快速、输出概率', '只能线性可分'),
    ('决策树', '分类/回归', '可解释、非线性', '极易过拟合'),
    ('随机森林', '分类/回归', '准确率高、抗过拟合', '速度慢、黑盒'),
    ('KNN', '分类/回归', '简单、非线性', '预测慢'),
]
for row_data in overview_data:
    row = t_overview.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val
        set_cell_font(row[i])

add_paragraph_with_format(doc, '', first_line_indent=False)

add_heading_custom(doc, '2.2 线性回归（Linear Regression）')

add_paragraph_with_format(doc,
    '线性回归是最基础的机器学习算法，其目标是通过学习一个线性函数来拟合自变量与因变量之间的关系。'
    '模型形式为 y = w^T * x + b，其中 w 为权重向量，b 为偏置项。'
    '线性回归通过最小化均方误差（MSE）来求解参数，可以使用正规方程或梯度下降法进行优化。'
    '均方误差的数学形式为：MSE = (1/n) * sum((y_i - y_hat_i)^2)，'
    '目标是找到使 MSE 最小的 w 和 b。')

add_paragraph_with_format(doc,
    '虽然线性回归本质上是一种回归算法，但也可以用于二分类任务。'
    '具体做法是：将类别标签编码为0和1，使用线性回归拟合后，'
    '将预测值以0.5为阈值进行划分——预测值大于等于0.5判定为正类（1），小于0.5判定为负类（0）。'
    '同时，线性回归的连续输出值可以作为样本属于正类的"评分"，用于绘制ROC曲线和计算AUC。'
    '这种做法的局限性在于：线性回归的输出不受限于(0,1)区间，'
    '对于远离决策边界的样本可能给出极端预测值，且不满足概率的统计性质。'
    '正是为了解决这一问题，逻辑回归在输出端引入了Sigmoid函数，将线性组合的结果映射到概率空间。'
    '线性回归的优点是训练速度极快、可解释性最强（权重直接反映特征影响方向和大小），'
    '在金融领域常用于收益率预测、因子打分和资产定价模型。')

add_heading_custom(doc, '2.3 逻辑回归（Logistic Regression）')

add_paragraph_with_format(doc,
    '逻辑回归虽然名字中带有"回归"，但实际上是一种经典的二分类算法。'
    '其核心思想是在线性回归的基础上，通过Sigmoid函数将线性组合的输出映射到(0, 1)区间，'
    '作为样本属于正类的概率。Sigmoid函数的数学形式为：sigma(z) = 1 / (1 + e^(-z))，'
    '其中 z = w^T * x + b 是特征的线性组合。当概率大于阈值（通常为0.5）时，'
    '将样本预测为正类，否则预测为负类。')

add_paragraph_with_format(doc,
    '逻辑回归的损失函数采用交叉熵损失（Cross-Entropy Loss），'
    '通过梯度下降法优化模型参数。其决策边界是线性的，即 w^T * x + b = 0 定义的超平面。'
    '逻辑回归的优点在于训练速度快、可解释性强（系数可解读为特征对预测概率的影响），'
    '且天然输出概率值，支持灵活的阈值调整。'
    '缺点是只能学习线性决策边界，对非线性关系需要手动构造交互特征或多项式特征。'
    '在金融领域，逻辑回归常用于信用评分、违约预测和涨跌方向预测的baseline模型。')

add_heading_custom(doc, '2.4 决策树（Decision Tree）')

add_paragraph_with_format(doc,
    '决策树是一种基于树形结构进行决策的算法。它通过递归地选择最优特征和切分点，'
    '将特征空间划分为若干互不重叠的区域，每个叶子节点对应一个预测结果。'
    '决策树的构建过程是一个自顶向下的递归过程：在每个节点处，算法遍历所有特征和所有可能的切分点，'
    '选择使划分后纯度提升最大的特征作为当前节点的分裂特征。')

add_paragraph_with_format(doc,
    '分类决策树常用的划分准则包括基尼系数（Gini Index）和信息增益（Information Gain）。'
    '基尼系数衡量节点中样本类别的不纯度，值越小表示纯度越高；'
    '信息增益基于信息论中的熵概念，衡量划分前后不确定性的减少量。'
    '决策树的关键超参数包括最大深度（max_depth）、节点最小样本数（min_samples_split）'
    '和叶子节点最小样本数（min_samples_leaf），这些参数用于控制树的复杂度，防止过拟合。')

add_paragraph_with_format(doc,
    '决策树的优点是可解释性强（可以直接画出树结构，理解每条决策路径）、'
    '能处理非线性关系、对特征量纲不敏感（无需标准化）。'
    '缺点是极易过拟合，单棵树的方差很大，对训练数据的微小变化敏感。'
    '在金融领域，决策树可用于规则化交易策略生成和特征重要性排序。')

add_heading_custom(doc, '2.5 随机森林（Random Forest）')

add_paragraph_with_format(doc,
    '随机森林是一种基于Bagging思想的集成学习算法。它通过同时训练多棵决策树，'
    '并对所有树的预测结果进行投票（分类）或平均（回归），从而获得比单棵决策树更优的性能。'
    '随机森林的核心创新在于引入了两层随机性：'
    '第一层是样本随机，每棵树使用Bootstrap抽样（有放回抽样）从训练集中获取一个子集；'
    '第二层是特征随机，每个节点分裂时只考虑随机选取的特征子集，而非全部特征。'
    '这两层随机性使得各棵树之间具有足够的多样性，从而有效降低了集成模型的方差。')

add_paragraph_with_format(doc,
    '随机森林的优点包括：准确率通常显著高于单棵决策树、抗过拟合能力强'
    '（多棵树的平均抵消了单棵树的过拟合）、自带特征重要性评估'
    '（基于不纯度减少量计算）、支持并行训练。'
    '缺点是训练速度较慢（需要训练大量树）、模型可解释性较差（黑盒模型）、模型体积较大。'
    '在金融领域，随机森林常用于涨跌方向预测、风险因子识别和特征筛选。')

add_heading_custom(doc, '2.6 K近邻（KNN）')

add_paragraph_with_format(doc,
    'K近邻（K-Nearest Neighbors, KNN）是一种基于实例的懒惰学习算法。'
    '与上述算法不同，KNN在训练阶段不显式地学习模型参数，'
    '而是将训练数据全部存储，在预测阶段通过计算待预测样本与训练样本之间的距离，'
    '找到最近的K个邻居，然后根据这K个邻居的标签进行投票（分类）或平均（回归）。'
    '常用的距离度量包括欧氏距离、曼哈顿距离和闵可夫斯基距离。')

add_paragraph_with_format(doc,
    'KNN的核心超参数是K值（邻居数量）。K值较小意味着模型对噪声敏感、容易过拟合；'
    'K值较大则会使决策边界过于平滑、可能欠拟合。'
    '通常通过交叉验证选择最优K值。KNN的优点是原理简单直观、无需训练过程（适合快速原型验证）、'
    '天然支持非线性决策边界。'
    '缺点是预测速度慢（每次预测都要计算与所有训练样本的距离）、'
    '对特征量纲敏感（必须标准化）、在高维空间中距离度量失效（维度灾难）。'
    '在金融领域，KNN可用于相似股票筛选和模式匹配交易。')

add_heading_custom(doc, '2.7 五种算法对比小结')

# 算法对比表
add_paragraph_with_format(doc, '表2 五种算法多维度对比', first_line_indent=False)
t_cmp = doc.add_table(rows=1, cols=6)
t_cmp.style = 'Table Grid'
hdr_cmp = t_cmp.rows[0].cells
headers_cmp = ['维度', '线性回归', '逻辑回归', '决策树', '随机森林', 'KNN']
for i, h in enumerate(headers_cmp):
    hdr_cmp[i].text = h
    set_cell_font(hdr_cmp[i], bold=True)

cmp_data = [
    ('训练速度', '最快', '快', '中等', '较慢', '无需训练'),
    ('预测速度', '快', '快', '快', '中等', '慢'),
    ('可解释性', '高（权重可读）', '高（系数可读）', '高（树形可视化）', '低（黑盒）', '低'),
    ('准确率', '中等', '中等', '中等', '高', '中等'),
    ('过拟合风险', '低', '低', '高', '低', '中等'),
    ('非线性能力', '弱', '弱', '强', '强', '强'),
    ('输出类型', '连续值', '概率', '类别/概率', '类别/概率', '类别/概率'),
    ('需标准化', '是', '是', '否', '否', '是'),
]
for row_data in cmp_data:
    row = t_cmp.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val
        set_cell_font(row[i])

add_paragraph_with_format(doc,
    '从对比可以看出，五种算法各有优劣。在实际建模中，建议遵循"先简后繁"的原则：'
    '先用线性回归或逻辑回归建立baseline，了解数据的基本可分性；'
    '再尝试决策树和KNN探索非线性特征；最后使用随机森林追求最高准确率。'
    '同时，应结合具体任务的需求（如是否需要可解释性、是否需要概率输出）选择合适的算法。')

add_paragraph_with_format(doc, '', first_line_indent=False)

# ==================== 三、机器学习模型评价指标 ====================
add_heading_custom(doc, '三、机器学习模型评价指标')

add_heading_custom(doc, '3.1 数据集划分原则')

add_paragraph_with_format(doc,
    '在构建机器学习模型之前，首先需要将数据集划分为训练集和测试集。'
    '训练集用于模型学习参数，测试集用于评估模型的泛化能力。'
    '对于非时间序列数据（如医学诊断数据），可以使用随机划分的方式，'
    '即通过train_test_split函数随机抽取一定比例的样本作为测试集。'
    '随机划分时应注意使用stratify参数保持类别比例一致。')

add_paragraph_with_format(doc,
    '对于时间序列数据（如股票行情数据），则不能使用随机划分，必须按照时间顺序进行切分。'
    '这是因为时间序列数据具有时间依赖性，如果随机打乱，会导致训练集中包含测试集时间点之后的数据，'
    '即"用未来数据预测过去"，这种数据泄露会使模型评估结果严重虚高，'
    '在实盘中完全无法复现。正确的做法是将时间序列按时间排列，'
    '前段作为训练集，后段作为测试集，确保模型始终只用历史数据预测未来。'
    '本任务中，乳腺癌数据集采用随机划分，股票数据采用时间顺序划分，'
    '通过对比两种划分方式加深对这一原则的理解。')

add_heading_custom(doc, '3.2 混淆矩阵')

add_paragraph_with_format(doc,
    '混淆矩阵是评估分类模型性能的基础工具，它以矩阵形式展示模型预测结果与真实标签之间的对照关系。'
    '对于二分类问题，混淆矩阵是一个2x2的表格，包含四个基本元素：'
    '真阳性（TP，True Positive）表示实际为正且预测为正的样本数；'
    '真阴性（TN，True Negative）表示实际为负且预测为负的样本数；'
    '假阳性（FP，False Positive）表示实际为负但预测为正的样本数（又称第一类错误）；'
    '假阴性（FN，False Negative）表示实际为正但预测为负的样本数（又称第二类错误）。')

add_paragraph_with_format(doc,
    '在金融场景中，如果将"次日上涨"定义为正类（1），"次日下跌"定义为负类（0），'
    '则混淆矩阵的四个元素具有明确的金融含义：'
    'TP表示模型正确预测上涨并实际也上涨（正确买入获利）；'
    'TN表示模型正确预测下跌并实际也下跌（正确回避避免亏损）；'
    'FP表示模型预测上涨但实际下跌（错误买入导致亏损，这是最危险的错误）；'
    'FN表示模型预测下跌但实际上涨（错过盈利机会，相对可接受）。'
    '因此，在金融分类任务中，控制FP（假阳性率）通常比控制FN更重要。')

add_paragraph_with_format(doc,
    '基于混淆矩阵的四个基本元素，可以衍生出多个评估指标：'
    '准确率（Accuracy）= (TP+TN) / (TP+TN+FP+FN)，表示整体预测正确率，'
    '但在类别不均衡时会失真；'
    '精确率（Precision）= TP / (TP+FP)，表示预测为正的样本中实际为正的比例，'
    '关注误报率；'
    '召回率（Recall）= TP / (TP+FN)，表示实际为正的样本中被正确预测的比例，'
    '关注漏报率；'
    'F1分数 = 2 * Precision * Recall / (Precision + Recall)，是精确率和召回率的调和平均，'
    '在两者之间取得平衡。')

add_heading_custom(doc, '3.3 ROC曲线与AUC')

add_paragraph_with_format(doc,
    'ROC曲线（Receiver Operating Characteristic Curve，受试者工作特征曲线）'
    '是评估二分类模型性能的重要工具。ROC曲线的横轴为假阳性率（FPR = FP / (FP+TN)），'
    '纵轴为真阳性率（TPR = TP / (TP+FN)，即召回率）。'
    'ROC曲线的绘制过程如下：首先，模型对每个测试样本输出一个属于正类的概率值；'
    '然后，设定一个阈值（threshold），将概率大于等于阈值的样本预测为正类，小于阈值的预测为负类；'
    '在该阈值下计算(FPR, TPR)作为一个点；'
    '最后，遍历所有可能的阈值（从0到1），将所有点连接成一条曲线，即为ROC曲线。')

add_paragraph_with_format(doc,
    'ROC曲线的关键性质包括：曲线越靠近左上角，模型性能越好；'
    '对角线（从(0,0)到(1,1)的直线）代表随机猜测，其AUC为0.5；'
    '如果ROC曲线低于对角线，说明模型预测方向与真实情况相反，比随机猜测还差。'
    'ROC曲线的优点是对类别不均衡稳健，且不依赖于具体阈值的选择，'
    '能够全面反映模型在不同阈值下的整体排序能力。')

add_paragraph_with_format(doc,
    'AUC（Area Under Curve）即ROC曲线下方的面积，取值范围为[0, 1]。'
    'AUC的概率含义是：随机抽取一个正样本和一个负样本，模型将正样本排在负样本前面的概率。'
    'AUC的评级标准通常为：AUC等于1.0表示完美分类；AUC大于0.8表示模型良好；'
    'AUC大于0.7表示模型有效；AUC等于0.5表示模型无区分能力（等同于随机猜测）；'
    'AUC小于0.5表示模型预测方向错误。'
    '与准确率相比，AUC不受类别分布和阈值选择的影响，'
    '是衡量模型排序能力的综合指标，在金融预测任务中被广泛采用。')

add_paragraph_with_format(doc, '', first_line_indent=False)

# ==================== 四、实验与结果 ====================
add_heading_custom(doc, '四、实验与结果')

add_heading_custom(doc, '4.1 数据集说明')

add_heading_custom(doc, '4.1.1 乳腺癌数据集')

add_paragraph_with_format(doc,
    '本任务使用scikit-learn内置的乳腺癌诊断数据集（Breast Cancer Wisconsin Diagnostic Dataset）'
    '作为通用二分类基准实验。该数据集包含569个样本，每个样本有30个数值型特征，'
    '包括细胞核的半径、纹理、周长、面积、平滑度、紧密度、凹度、凹点、对称性和分形维度等'
    '十个基本特征的均值、标准差和最大值。目标变量为二分类标签：'
    '0表示恶性（Malignant），1表示良性（Benign）。'
    '数据集中良性样本357个，恶性样本212个，类别分布基本均衡。')

add_heading_custom(doc, '4.1.2 股票财务指标收益数据')

add_paragraph_with_format(doc,
    '本任务复用TASK1中已下载的长江电力（600900.SH）日线数据作为金融场景实验数据。'
    '原始数据包含2025年7月2日至2026年7月2日共243个交易日的开盘价、最高价、最低价、'
    '收盘价、前收盘价、涨跌幅和成交量等字段。在原始数据基础上，'
    '通过特征工程构造了11个技术指标特征：'
    'MA5偏离度、MA10偏离度、MA20偏离度（收盘价相对移动平均线的偏离程度）、'
    'RSI(14)（相对强弱指标，衡量超买超卖程度）、'
    '成交量变化率、日内波动率、当日涨跌幅、5日动量、10日动量、'
    '5日波动率和量比（当日成交量与5日均量之比）。'
    '标签定义为次日涨跌方向：1表示次日上涨（涨跌幅大于0），0表示次日下跌。'
    '经过去除NaN行后，有效样本数为224个，其中正类（次日涨）101个，负类（次日跌）123个。')

add_paragraph_with_format(doc, '', first_line_indent=False)

add_heading_custom(doc, '4.2 数据划分与预处理')

add_paragraph_with_format(doc,
    '乳腺癌数据集采用随机划分方式，使用train_test_split函数按80:20的比例划分训练集和测试集，'
    '并设置stratify=y保持训练集和测试集中的类别比例一致，随机种子固定为42以确保结果可复现。'
    '划分后训练集455个样本，测试集114个样本。')

add_paragraph_with_format(doc,
    '股票数据集采用时间顺序划分方式，严格按照时间先后顺序，'
    '前80%的交易日（179个样本）作为训练集，后20%的交易日（45个样本）作为测试集。'
    '这种划分方式确保模型只用历史数据训练，用未来数据测试，避免了数据泄露问题。'
    '需要特别强调的是，金融时间序列数据绝不能使用随机打乱划分，'
    '否则会引入"未来函数"导致评估结果虚高。')

add_paragraph_with_format(doc,
    '在特征预处理方面，两个数据集均使用StandardScaler进行标准化处理'
    '（均值为0，标准差为1）。标准化过程严格遵循"在训练集上fit，在测试集上transform"的原则，'
    '避免测试集信息泄露到训练过程中。虽然决策树和随机森林对特征量纲不敏感，'
    '但线性回归、逻辑回归和KNN对量纲敏感，因此为了多模型公平对比，统一进行标准化处理。')

add_paragraph_with_format(doc, '', first_line_indent=False)

add_heading_custom(doc, '4.3 模型构建与训练')

add_paragraph_with_format(doc,
    '本任务构建了五种模型进行对比实验：'
    '线性回归（Linear Regression），使用最小二乘法拟合，预测值以0.5为阈值进行分类，'
    '将连续输出裁剪到[0,1]区间作为评分用于ROC/AUC计算；'
    '逻辑回归（Logistic Regression），设置max_iter=1000确保收敛，作为分类baseline模型；'
    '决策树（Decision Tree），设置max_depth=5控制树的复杂度，防止过拟合；'
    '随机森林（Random Forest），设置n_estimators=100（100棵决策树）、max_depth=10，'
    '启用并行训练（n_jobs=-1），作为主要预测模型；'
    'KNN（K-Nearest Neighbors），设置n_neighbors=5，使用欧氏距离度量。'
    '五种模型均在相同的训练集上训练，在相同的测试集上评估，确保对比的公平性。')

add_paragraph_with_format(doc, '', first_line_indent=False)

# ==================== 4.4 乳腺癌数据集实验结果 ====================
add_heading_custom(doc, '4.4 乳腺癌数据集实验结果')

add_paragraph_with_format(doc,
    '在乳腺癌数据集上，五种模型均取得了优秀的分类效果。'
    '下表展示了各模型的详细评估指标。')

add_results_table(doc, cancer_results, '表3 乳腺癌数据集五种模型评估结果')

add_paragraph_with_format(doc,
    '图1展示了决策树模型在乳腺癌测试集上的混淆矩阵。'
    '在114个测试样本中，决策树正确分类了105个样本（TP=66，TN=39），'
    '错误分类了9个样本（FP=3，FN=6）。整体准确率为92.11%，'
    '但假阴性（FN=6）相对较多，意味着有6个恶性样本被误判为良性，'
    '在医学诊断场景中这类错误的风险较高。')

img1 = os.path.join(IMAGE_DIR, 'breast_cancer_fig1_cm_dt.png')
if os.path.exists(img1):
    add_image_with_caption(doc, img1, '图1 乳腺癌数据集决策树模型混淆矩阵')

add_paragraph_with_format(doc,
    '图2展示了随机森林模型的混淆矩阵。相比决策树，随机森林的正确分类数提升至109个'
    '（TP=70，TN=39），错误分类降至5个（FP=3，FN=2）。'
    '假阴性从6降至2，说明随机森林在减少漏诊方面优于单棵决策树。'
    '这体现了集成学习通过多棵树投票降低方差、提升稳健性的优势。')

img2 = os.path.join(IMAGE_DIR, 'breast_cancer_fig2_cm_rf.png')
if os.path.exists(img2):
    add_image_with_caption(doc, img2, '图2 乳腺癌数据集随机森林模型混淆矩阵')

add_paragraph_with_format(doc,
    '图3展示了线性回归模型的混淆矩阵。线性回归正确分类了109个样本'
    '（TP=71，TN=38），错误分类6个（FP=4，FN=1）。'
    '值得注意的是，线性回归的假阴性仅为1，在所有模型中最低，'
    '说明其几乎不会漏掉良性样本。但假阳性为4，略高于逻辑回归的1，'
    '表明线性回归在没有Sigmoid约束的情况下，对少数边界样本的判别稍显粗糙。'
    '尽管如此，其AUC仍达到0.9924，说明线性回归在乳腺癌这种特征-目标关系较为线性的数据集上同样有效。')

img3 = os.path.join(IMAGE_DIR, 'breast_cancer_fig3_cm_lr.png')
if os.path.exists(img3):
    add_image_with_caption(doc, img3, '图3 乳腺癌数据集线性回归模型混淆矩阵')

add_paragraph_with_format(doc,
    '图4展示了KNN模型的混淆矩阵。KNN正确分类了109个样本'
    '（TP=70，TN=39），错误分类5个（FP=3，FN=2），'
    '与随机森林的表现完全一致。这说明在经过标准化的乳腺癌数据上，'
    '基于距离的KNN能够有效识别样本的类别聚类结构。'
    'KNN的AUC为0.9788，略低于逻辑回归和随机森林，'
    '这可能是因为KNN的决策边界是局部的、不连续的，'
    '在概率排序的精细程度上不如全局优化的线性模型和集成模型。')

img4 = os.path.join(IMAGE_DIR, 'breast_cancer_fig4_cm_knn.png')
if os.path.exists(img4):
    add_image_with_caption(doc, img4, '图4 乳腺癌数据集KNN模型混淆矩阵')

add_paragraph_with_format(doc, '', first_line_indent=False)

add_paragraph_with_format(doc,
    '图5展示了五种模型在乳腺癌数据集上的ROC曲线对比。'
    '五条ROC曲线均显著偏离对角线，向左上角弯曲，说明五种模型均具有较强的分类能力。'
    '其中逻辑回归的AUC最高（0.9954），随机森林次之（0.9939），线性回归第三（0.9924），'
    'KNN第四（0.9788），决策树最低（0.9163）。'
    '逻辑回归和线性回归在该数据集上表现优异，可能是因为乳腺癌数据集的特征与目标之间'
    '存在较强的线性关系，线性模型的决策边界已经能够很好地捕捉这种关系。'
    '决策树的AUC相对较低，这是因为单棵决策树容易过拟合，泛化能力不如集成模型和线性模型。')

img5 = os.path.join(IMAGE_DIR, 'breast_cancer_fig5_roc.png')
if os.path.exists(img5):
    add_image_with_caption(doc, img5, '图5 乳腺癌数据集五种模型ROC曲线对比')

add_paragraph_with_format(doc, '', first_line_indent=False)

add_paragraph_with_format(doc,
    '图6以柱状图形式直观对比了五种模型的AUC值。'
    '逻辑回归、随机森林和线性回归的AUC均接近1.0，差异很小；'
    'KNN的AUC为0.9788，也处于优秀水平；'
    '决策树的AUC为0.9163，与前四者有一定差距。'
    '整体来看，在结构化的医学诊断数据上，机器学习模型能够达到接近完美的分类效果，'
    '且模型之间的差异不大，说明数据本身具有良好的可分性。')

img6 = os.path.join(IMAGE_DIR, 'breast_cancer_fig6_auc_bar.png')
if os.path.exists(img6):
    add_image_with_caption(doc, img6, '图6 乳腺癌数据集五种模型AUC对比')

add_paragraph_with_format(doc, '', first_line_indent=False)

add_paragraph_with_format(doc,
    '图7展示了随机森林模型的特征重要性排序。'
    '特征重要性反映了各特征对模型预测的贡献程度，由随机森林在训练过程中自动计算。'
    '从图中可以看出，最重要的特征主要集中在细胞核的形态学指标上，'
    '如最差凹点数（worst concave points）、最差半径（worst radius）等，'
    '这与医学上恶性肿瘤细胞核形态不规则、大小异常的常识一致。'
    '特征重要性分析不仅有助于理解模型的决策依据，还能指导特征筛选和降维。')

img7 = os.path.join(IMAGE_DIR, 'breast_cancer_fig7_feature_imp.png')
if os.path.exists(img7):
    add_image_with_caption(doc, img7, '图7 乳腺癌数据集随机森林特征重要性')

add_paragraph_with_format(doc, '', first_line_indent=False)

# ==================== 4.5 股票数据集实验结果 ====================
add_heading_custom(doc, '4.5 股票数据集实验结果')

add_paragraph_with_format(doc,
    '在长江电力股票数据上，五种模型的表现与乳腺癌数据集形成了鲜明对比。'
    '下表展示了各模型的评估指标。')

add_results_table(doc, stock_results, '表4 股票数据集五种模型评估结果')

add_paragraph_with_format(doc,
    '从表4可以看出，五种模型在股票数据上的准确率均在38%~47%之间，'
    '不仅低于50%的随机基线，AUC值也均低于0.5'
    '（线性回归0.344、逻辑回归0.346、决策树0.459、随机森林0.391、KNN 0.360）。'
    'AUC低于0.5意味着模型的预测方向与真实涨跌方向呈反向关系，'
    '即模型预测"涨"的样本反而更可能"跌"。'
    '这一结果虽然不理想，但却非常真实地反映了金融预测的困难性。')

add_paragraph_with_format(doc,
    '图8展示了决策树模型在股票测试集上的混淆矩阵。'
    '在45个测试样本中，决策树仅正确分类了18个样本（TP=9，TN=9），'
    '错误分类27个样本（FP=14，FN=13），准确率仅为40%。'
    '特别是FP=14，意味着有14个实际下跌的交易日被模型预测为上涨，'
    '如果据此进行交易，将直接导致14次亏损操作。')

img8 = os.path.join(IMAGE_DIR, 'stock_600900_fig1_cm_dt.png')
if os.path.exists(img8):
    add_image_with_caption(doc, img8, '图8 股票数据集决策树模型混淆矩阵')

add_paragraph_with_format(doc,
    '图9展示了随机森林模型的混淆矩阵。随机森林的正确分类数提升至21个'
    '（TP=8，TN=13），准确率46.67%，略优于决策树，但仍然低于随机基线。'
    '随机森林在识别下跌样本（TN=13）方面表现稍好，但在识别上涨样本（TP=8）方面依然较弱。')

img9 = os.path.join(IMAGE_DIR, 'stock_600900_fig2_cm_rf.png')
if os.path.exists(img9):
    add_image_with_caption(doc, img9, '图9 股票数据集随机森林模型混淆矩阵')

add_paragraph_with_format(doc,
    '图10展示了线性回归模型的混淆矩阵。线性回归正确分类了19个样本'
    '（TP=5，TN=14），准确率42.22%。'
    '线性回归在识别下跌样本（TN=14）方面是五种模型中最多的，'
    '但在识别上涨样本（TP=5）方面表现最弱，'
    '说明线性回归在股票数据上倾向于将样本预测为负类（下跌）。'
    '这可能与线性回归的输出不受Sigmoid约束有关，'
    '在训练集正负类不均衡的情况下（负类123个 vs 正类101个），'
    '线性回归的决策阈值偏离了最优位置。')

img10 = os.path.join(IMAGE_DIR, 'stock_600900_fig3_cm_lr.png')
if os.path.exists(img10):
    add_image_with_caption(doc, img10, '图10 股票数据集线性回归模型混淆矩阵')

add_paragraph_with_format(doc,
    '图11展示了KNN模型的混淆矩阵。KNN仅正确分类了17个样本'
    '（TP=6，TN=11），准确率37.78%，是五种模型中最低的。'
    'KNN在股票数据上表现最差，可能原因有二：'
    '一是股票数据的11维特征空间中，样本之间的距离区分度不高，KNN难以找到有意义的近邻；'
    '二是训练集仅179个样本，在高维空间中过于稀疏，KNN的"维度灾难"效应明显。')

img11 = os.path.join(IMAGE_DIR, 'stock_600900_fig4_cm_knn.png')
if os.path.exists(img11):
    add_image_with_caption(doc, img11, '图11 股票数据集KNN模型混淆矩阵')

add_paragraph_with_format(doc, '', first_line_indent=False)

add_paragraph_with_format(doc,
    '图12展示了五种模型在股票数据集上的ROC曲线对比。'
    '与乳腺癌数据集形成鲜明对比的是，五条ROC曲线均位于对角线下方或紧贴对角线，'
    'AUC值均低于0.5，说明模型的排序能力为负向，'
    '即模型倾向于给实际下跌的样本更高的上涨概率。'
    '其中决策树的AUC（0.459）最接近0.5，但也未达到有效预测的水平。')

img12 = os.path.join(IMAGE_DIR, 'stock_600900_fig5_roc.png')
if os.path.exists(img12):
    add_image_with_caption(doc, img12, '图12 股票数据集五种模型ROC曲线对比')

add_paragraph_with_format(doc, '', first_line_indent=False)

add_paragraph_with_format(doc,
    '这一现象的可能原因包括以下几个方面。'
    '第一，样本量过小。224个有效样本中训练集仅179个，测试集仅45个，'
    '对于11维特征的模型而言，数据量严重不足，模型难以学习到稳定的模式。'
    '第二，金融市场具有弱有效性。根据有效市场假说，'
    '当前价格已经反映了所有历史信息，基于历史技术指标预测次日涨跌本身就极具挑战性。'
    '第三，训练集与测试集跨越不同的市场阶段。'
    '训练集覆盖了前80%的交易日，测试集覆盖了最后20%的交易日，'
    '两个时期的市场环境可能存在显著差异（如训练期偏上涨，测试期偏下跌），'
    '导致模型在训练集上学到的模式在测试集上失效甚至反向。'
    '第四，简单技术指标的信息含量有限。本任务使用的MA偏离度、RSI等技术指标'
    '是市场广泛使用的基础指标，其包含的预测信息可能已被市场充分定价。')

add_paragraph_with_format(doc,
    '图13以柱状图形式对比了五种模型在股票数据上的AUC值。'
    '五种模型的AUC均低于0.5的随机基线，其中决策树的AUC（0.459）最接近0.5，'
    'KNN（0.360）和线性回归（0.344）最低。'
    '这一结果在量化交易领域具有普遍性参考意义：'
    '简单的机器学习模型加上基础技术指标，难以在金融市场中获得稳定的预测优势。'
    '要提升预测效果，需要更丰富的特征工程（如基本面因子、情绪因子、'
    '宏观经济因子等）、更大的数据量、以及更复杂的模型架构。')

img13 = os.path.join(IMAGE_DIR, 'stock_600900_fig6_auc_bar.png')
if os.path.exists(img13):
    add_image_with_caption(doc, img13, '图13 股票数据集五种模型AUC对比')

add_paragraph_with_format(doc, '', first_line_indent=False)

add_paragraph_with_format(doc,
    '图14展示了随机森林在股票数据上的特征重要性排序。'
    '从图中可以看出，RSI(14)、10日动量和MA5偏离度是最重要的三个特征，'
    '但所有特征的重要性分布相对均匀，没有某个特征占据绝对主导地位。'
    '这说明在股票预测任务中，单个技术指标的预测力有限，'
    '模型难以找到强信号特征，这也是AUC偏低的原因之一。')

img14 = os.path.join(IMAGE_DIR, 'stock_600900_fig7_feature_imp.png')
if os.path.exists(img14):
    add_image_with_caption(doc, img14, '图14 股票数据集随机森林特征重要性')

add_paragraph_with_format(doc, '', first_line_indent=False)

# ==================== 4.6 实验结论 ====================
add_heading_custom(doc, '4.6 实验结论')

add_paragraph_with_format(doc,
    '通过两个数据集的对比实验，可以得出以下结论：'
    '第一，机器学习模型在结构化数据上表现优异。乳腺癌数据集上五种模型的AUC均超过0.91，'
    '逻辑回归和随机森林更是接近完美（AUC > 0.99），'
    '说明当特征与目标之间存在明确的映射关系时，机器学习能够有效捕捉这种关系。')

add_paragraph_with_format(doc,
    '第二，金融时间序列预测具有极高的挑战性。股票数据集上五种模型的AUC均低于0.5，'
    '远不如在医学数据上的表现。这一结果并非模型本身的问题，'
    '而是反映了金融市场的高度复杂性和有效性。'
    '简单技术指标和有限样本量难以支撑有效的涨跌预测，'
    '这也印证了量化交易中"没有银弹"的现实。')

add_paragraph_with_format(doc,
    '第三，时间序列数据的正确划分至关重要。本任务对股票数据采用时间顺序划分，'
    '虽然模型表现不佳，但这一结果是真实可信的。'
    '如果错误地使用随机划分，模型可能会因为数据泄露而获得虚高的评估指标，'
    '在实盘中将面临严重的亏损风险。')

add_paragraph_with_format(doc,
    '第四，模型选择应因地制宜。在乳腺癌数据集上，简单的逻辑回归反而表现最好，'
    '说明并非越复杂的模型越好；而在股票数据上，五种模型的表现差异不大，'
    '说明数据质量和特征工程才是决定性因素。'
    '实际应用中应先建立baseline，再逐步尝试更复杂的模型和更丰富的特征。')

add_paragraph_with_format(doc,
    '第五，线性回归与逻辑回归在分类任务中表现相近但有差异。'
    '在乳腺癌数据集上两者AUC相差仅0.003，但在股票数据上，'
    '线性回归由于缺乏Sigmoid约束，对阈值更敏感。'
    '这印证了对于分类任务，逻辑回归在理论上是比线性回归更合适的选择，'
    '但在特征-目标关系接近线性的场景下，线性回归同样具有竞争力。')

add_paragraph_with_format(doc, '', first_line_indent=False)

# ==================== 五、总结与展望 ====================
add_heading_custom(doc, '五、总结与展望')

add_heading_custom(doc, '5.1 本任务收获')

add_paragraph_with_format(doc,
    '通过本次TASK5的实践，我对机器学习算法有了系统的认识和理解。'
    '在理论层面，深入学习了线性回归、逻辑回归、决策树、随机森林和KNN五种算法的原理、优缺点和适用场景，'
    '掌握了混淆矩阵、AUC、ROC曲线等核心评价指标的计算方法和解读能力。'
    '在实践层面，通过Python编程完成了从数据加载、特征工程、数据划分、'
    '模型训练到模型评估的完整机器学习流程，并在两个不同类型的数据集上进行了对比实验。')

add_paragraph_with_format(doc,
    '最深刻的收获来自于两个数据集的对比实验结果。'
    '乳腺癌数据集上模型表现优异，AUC接近1.0，展示了机器学习在结构化数据上的强大能力；'
    '而股票数据集上模型AUC低于0.5，真实地揭示了金融预测的困难性。'
    '这一对比让我深刻认识到，机器学习并非万能工具，'
    '其效果在很大程度上取决于数据质量和问题本身的可预测性。'
    '在量化交易中应用机器学习时，必须对市场的有效性保持敬畏，'
    '不能盲目依赖模型的预测结果。')

add_heading_custom(doc, '5.2 局限与改进方向')

add_paragraph_with_format(doc,
    '本任务存在以下局限性：'
    '第一，股票数据样本量较小（224个有效样本），不足以支撑复杂的机器学习模型；'
    '第二，特征工程较为基础，仅使用了MA偏离度、RSI等技术指标，'
    '未包含基本面因子、情绪因子等更丰富的信息源；'
    '第三，模型超参数未经系统调优，可能存在提升空间；'
    '第四，未考虑交易成本、滑点等实际交易因素。')

add_paragraph_with_format(doc,
    '未来的改进方向包括：'
    '第一，扩大数据范围，使用更多股票、更长时间跨度的数据；'
    '第二，丰富特征工程，引入基本面因子（如市盈率、市净率）、'
    '情绪因子（如新闻情感分析）、宏观经济因子（如利率、CPI）等；'
    '第三，使用更先进的模型架构，如XGBoost、LightGBM、神经网络等；'
    '第四，进行系统的超参数调优，如网格搜索或贝叶斯优化；'
    '第五，将分类预测与具体的交易策略结合，'
    '在回测框架中评估模型的实际盈利能力，而非仅看AUC指标。')

# ==================== 保存文档 ====================
output_dir = r'C:\Users\LENOVO\Desktop\quant-ai\TASK5'
docx_path = os.path.join(output_dir, '林富强+TASK5.docx')
doc.save(docx_path)
print(f"Word文档已保存至: {docx_path}")

# ==================== 转换为PDF ====================
try:
    from docx2pdf import convert
    pdf_path = os.path.join(output_dir, '林富强+TASK5.pdf')
    convert(docx_path, pdf_path)
    print(f"PDF文档已保存至: {pdf_path}")
except Exception as e:
    print(f"PDF转换失败: {e}")
    print("请手动将docx转换为PDF，或检查Microsoft Word是否已安装。")

# -*- coding: utf-8 -*-
"""
TASK8 综合学习报告生成脚本
- 生成全部中文图表（统一风格、红涨绿跌、无英文变量名）
- 用 python-docx 构建 Word（宋体/五号/1.5倍行距/0段距/两端对齐）
- 用 Word COM (win32com) 直接转 PDF
作者：林富强 / 北京大学
"""
import os
import json
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib import font_manager
from matplotlib.patches import Rectangle

# ---------- 路径 ----------
ROOT = r"C:\Users\LENOVO\Desktop\quant-ai\TASK8"
IMG = os.path.join(ROOT, "images")
os.makedirs(IMG, exist_ok=True)
DOCX = os.path.join(ROOT, "林富强+TASK8.docx")
PDF = os.path.join(ROOT, "林富强+TASK8.pdf")

# ---------- 中文字体 ----------
FONT_PATH = r"C:\Windows\Fonts\simhei.ttf"
font_manager.fontManager.addfont(FONT_PATH)
plt.rcParams['font.sans-serif'] = ['SimHei']
plt.rcParams['axes.unicode_minus'] = False

# ---------- 配色（A股习惯：红涨绿跌） ----------
RED = '#C0392B'
GREEN = '#27AE60'
NAVY = '#2C3E50'
LBLUE = '#5DADE2'
GRAY = '#95A5A6'
ORANGE = '#E67E22'

def cjk():
    return font_manager.FontProperties(fname=FONT_PATH)

def save(fig, name):
    path = os.path.join(IMG, name)
    fig.savefig(path, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return path

def color_by_sign(vals):
    return [RED if v >= 0 else GREEN for v in vals]

# ======================================================================
# 图1：学习路径与任务体系（横向流程）
# ======================================================================
def fig_path():
    tasks = [('任务一', '数据引擎'), ('任务二', '指标构造'), ('任务三', '均线策略'),
             ('任务四', '海龟策略'), ('任务五', '机器学习'), ('任务六', '截面选股'),
             ('任务七', '实盘模拟')]
    n = len(tasks)
    fig, ax = plt.subplots(figsize=(9.2, 3.0))
    for i, (t, s) in enumerate(tasks):
        x = i * 1.28
        ax.add_patch(Rectangle((x, 0.45), 1.05, 0.85, fc=NAVY, ec='none', zorder=2))
        ax.text(x + 0.525, 0.92, t, ha='center', va='center', color='white',
                fontsize=11, fontweight='bold', fontproperties=cjk(), zorder=3)
        ax.text(x + 0.525, 0.62, s, ha='center', va='center', color='white',
                fontsize=9, fontproperties=cjk(), zorder=3)
        if i < n - 1:
            ax.annotate('', xy=(x + 1.18, 0.87), xytext=(x + 1.08, 0.87),
                        arrowprops=dict(arrowstyle='->', color=GRAY, lw=1.6))
    ax.set_xlim(-0.1, n * 1.28), ax.set_ylim(0.2, 1.5)
    ax.axis('off')
    ax.set_title('图1  七项任务的学习路径与能力演进', fontsize=12, fontweight='bold',
                 fontproperties=cjk(), color=NAVY, pad=10)
    return save(fig, 'fig1_path.png')

# ======================================================================
# 图2：均线交叉策略 不同参数组合 超额收益
# ======================================================================
def fig_ma_periods():
    labels = ['均线5/10', '均线5/15', '均线5/20', '均线10/20', '均线10/30']
    excess = [4.95, 3.98, 5.21, 5.35, 2.21]
    fig, ax = plt.subplots(figsize=(6.8, 4.0))
    bars = ax.bar(labels, excess, color=color_by_sign(excess), width=0.62, zorder=3)
    for b, v in zip(bars, excess):
        ax.text(b.get_x() + b.get_width() / 2, v + 0.08, f'{v:.2f}%',
                ha='center', va='bottom', fontsize=10, fontproperties=cjk())
    ax.axhline(0, color='#333333', lw=0.8)
    ax.set_ylabel('超额收益（%，相对买入持有）', fontsize=10, fontproperties=cjk())
    ax.set_title('图2  均线交叉策略不同参数组合的超额收益', fontsize=12, fontweight='bold',
                 fontproperties=cjk(), color=NAVY, pad=10)
    ax.grid(axis='y', color='#E5E5E5', zorder=0)
    ax.set_ylim(0, 6.2)
    return save(fig, 'fig2_ma_periods.png')

# ======================================================================
# 图3：海龟策略 三只标的 超额收益
# ======================================================================
def fig_turtle_stocks():
    labels = ['长江电力', '贵州茅台', '中国平安']
    excess = [-0.33, 9.81, 18.40]
    fig, ax = plt.subplots(figsize=(6.8, 4.0))
    bars = ax.bar(labels, excess, color=color_by_sign(excess), width=0.55, zorder=3)
    for b, v in zip(bars, excess):
        ax.text(b.get_x() + b.get_width() / 2, v + (0.4 if v >= 0 else -0.9),
                f'{v:+.2f}%', ha='center', va='bottom' if v >= 0 else 'top',
                fontsize=10, fontproperties=cjk())
    ax.axhline(0, color='#333333', lw=0.8)
    ax.set_ylabel('超额收益（%，相对买入持有）', fontsize=10, fontproperties=cjk())
    ax.set_title('图3  海龟策略在三只标的上的超额收益', fontsize=12, fontweight='bold',
                 fontproperties=cjk(), color=NAVY, pad=10)
    ax.grid(axis='y', color='#E5E5E5', zorder=0)
    ax.set_ylim(-2.5, 21)
    return save(fig, 'fig3_turtle_stocks.png')

# ======================================================================
# 图4：海龟策略 24组参数 超额收益热力图（红涨绿跌）
# ======================================================================
def fig_turtle_heatmap():
    stocks = ['长江电力', '贵州茅台', '中国平安']
    params = ['5/3', '5/5', '10/5', '10/10', '15/7', '20/10', '30/15', '55/20']
    data = [
        [0.92, 0.74, 8.48, 4.07, 5.25, -0.33, -1.89, -2.00],
        [1.80, -1.06, 10.62, 7.90, 9.53, 9.81, 10.34, 12.38],
        [8.69, 13.60, 13.53, 13.48, 15.50, 18.40, 15.11, 11.46],
    ]
    fig, ax = plt.subplots(figsize=(8.2, 3.4))
    im = ax.imshow(data, cmap='RdYlGn_r', aspect='auto', vmin=-5, vmax=20)
    ax.set_xticks(range(len(params)))
    ax.set_xticklabels(params, fontsize=9, fontproperties=cjk())
    ax.set_yticks(range(len(stocks)))
    ax.set_yticklabels(stocks, fontsize=10, fontproperties=cjk())
    for i in range(len(stocks)):
        for j in range(len(params)):
            v = data[i][j]
            ax.text(j, i, f'{v:+.1f}', ha='center', va='center',
                    fontsize=8.5, color='black', fontproperties=cjk())
    ax.set_title('图4  海龟策略24组参数的超额收益（红=跑赢，绿=跑输）', fontsize=12,
                 fontweight='bold', fontproperties=cjk(), color=NAVY, pad=10)
    cbar = fig.colorbar(im, ax=ax, fraction=0.035, pad=0.02)
    cbar.set_label('超额收益（%）', fontsize=9, fontproperties=cjk())
    return save(fig, 'fig4_turtle_heatmap.png')

# ======================================================================
# 图5：机器学习模型 两类数据集 AUC 对比
# ======================================================================
def fig_ml_auc():
    models = ['线性回归', '逻辑回归', '决策树', '随机森林', 'K近邻']
    auc_med = [0.9924, 0.9954, 0.9163, 0.9939, 0.9788]
    auc_stock = [0.3439, 0.3458, 0.4585, 0.3913, 0.3597]
    x = range(len(models))
    w = 0.36
    fig, ax = plt.subplots(figsize=(7.2, 4.2))
    b1 = ax.bar([i - w / 2 for i in x], auc_med, w, label='乳腺癌数据集', color=LBLUE, zorder=3)
    b2 = ax.bar([i + w / 2 for i in x], auc_stock, w, label='长江电力股票', color=ORANGE, zorder=3)
    for bars in (b1, b2):
        for b in bars:
            ax.text(b.get_x() + b.get_width() / 2, b.get_height() + 0.01,
                    f'{b.get_height():.2f}', ha='center', va='bottom', fontsize=8.5,
                    fontproperties=cjk())
    ax.axhline(0.5, color='#333333', ls='--', lw=1, zorder=2)
    ax.text(len(models) - 0.5, 0.52, '随机基准0.5', fontsize=8.5, color='#333333',
            fontproperties=cjk())
    ax.set_xticks(list(x))
    ax.set_xticklabels(models, fontsize=9.5, fontproperties=cjk())
    ax.set_ylabel('AUC（判别能力）', fontsize=10, fontproperties=cjk())
    ax.set_ylim(0, 1.15)
    ax.legend(prop=cjk(), fontsize=9, loc='upper right')
    ax.set_title('图5  机器学习模型在两类数据集上的判别能力对比', fontsize=12,
                 fontweight='bold', fontproperties=cjk(), color=NAVY, pad=10)
    ax.grid(axis='y', color='#E5E5E5', zorder=0)
    return save(fig, 'fig5_ml_auc.png')

# ======================================================================
# 图6：截面选股 四模型 年化收益与最大回撤
# ======================================================================
def fig_cs_return():
    models = ['线性回归', '决策树', '随机森林', '极端梯度提升', '等权基准']
    annual = [-3.83, 2.40, -1.30, 3.10, 2.25]
    fig, ax = plt.subplots(figsize=(7.2, 4.2))
    bars = ax.bar(models, annual, color=color_by_sign(annual), width=0.6, zorder=3)
    for b, v in zip(bars, annual):
        ax.text(b.get_x() + b.get_width() / 2, v + (0.1 if v >= 0 else -0.18),
                f'{v:+.2f}%', ha='center', va='bottom' if v >= 0 else 'top',
                fontsize=9.5, fontproperties=cjk())
    ax.axhline(0, color='#333333', lw=0.8)
    ax.set_ylabel('年化收益率（%）', fontsize=10, fontproperties=cjk())
    ax.set_title('图6  截面选股四种模型的年化收益率', fontsize=12, fontweight='bold',
                 fontproperties=cjk(), color=NAVY, pad=10)
    ax.grid(axis='y', color='#E5E5E5', zorder=0)
    ax.set_ylim(-5, 5)
    return save(fig, 'fig6_cs_return.png')

# ======================================================================
# 图7：截面选股 信息系数 IC 对比
# ======================================================================
def fig_cs_ic():
    models = ['线性回归', '决策树', '随机森林', '极端梯度提升']
    ic = [-0.0869, -0.0260, -0.0040, 0.0133]
    fig, ax = plt.subplots(figsize=(6.8, 4.0))
    bars = ax.bar(models, ic, color=color_by_sign(ic), width=0.58, zorder=3)
    for b, v in zip(bars, ic):
        ax.text(b.get_x() + b.get_width() / 2, v + (0.002 if v >= 0 else -0.004),
                f'{v:+.4f}', ha='center', va='bottom' if v >= 0 else 'top',
                fontsize=9.5, fontproperties=cjk())
    ax.axhline(0, color='#333333', lw=0.8)
    ax.set_ylabel('信息系数（IC）', fontsize=10, fontproperties=cjk())
    ax.set_title('图7  截面选股各模型的平均信息系数', fontsize=12, fontweight='bold',
                 fontproperties=cjk(), color=NAVY, pad=10)
    ax.grid(axis='y', color='#E5E5E5', zorder=0)
    ax.set_ylim(-0.11, 0.03)
    return save(fig, 'fig7_cs_ic.png')

# ======================================================================
# 图8：实战策略（TASK7）长区间与样本外表现
# ======================================================================
def fig_task7():
    labels = ['本策略·长区间', '沪深300·同期', '本策略·样本外']
    annual = [2.97, 19.58, -8.44]
    fig, ax = plt.subplots(figsize=(6.8, 4.0))
    bars = ax.bar(labels, annual, color=color_by_sign(annual), width=0.55, zorder=3)
    for b, v in zip(bars, annual):
        ax.text(b.get_x() + b.get_width() / 2, v + (0.4 if v >= 0 else -0.9),
                f'{v:+.2f}%', ha='center', va='bottom' if v >= 0 else 'top',
                fontsize=10, fontproperties=cjk())
    ax.axhline(0, color='#333333', lw=0.8)
    ax.set_ylabel('年化收益率（%）', fontsize=10, fontproperties=cjk())
    ax.set_title('图8  实战策略长区间与样本外表现对比', fontsize=12, fontweight='bold',
                 fontproperties=cjk(), color=NAVY, pad=10)
    ax.grid(axis='y', color='#E5E5E5', zorder=0)
    ax.set_ylim(-11, 22)
    return save(fig, 'fig8_task7.png')


# ======================================================================
# 构建文档
# ======================================================================
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SONG = '宋体'
HEI = '黑体'

def set_cell_font(cell, size=10.5, bold=False, align='center', font=SONG):
    for p in cell.paragraphs:
        p.alignment = {'center': WD_ALIGN_PARAGRAPH.CENTER,
                       'left': WD_ALIGN_PARAGRAPH.LEFT,
                       'right': WD_ALIGN_PARAGRAPH.RIGHT}[align]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        for r in p.runs:
            r.font.name = font
            r.font.size = Pt(size)
            r.font.bold = bold
            r.font.element.rPr.rFonts.set(qn('w:eastAsia'), font)

def add_table(doc, title, headers, rows, col_widths=None):
    # 表题
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_before = Pt(6)
    tp.paragraph_format.space_after = Pt(4)
    tr = tp.add_run(title)
    tr.font.name = SONG
    tr.font.size = Pt(10.5)
    tr.font.element.rPr.rFonts.set(qn('w:eastAsia'), SONG)
    tr.font.bold = True
    # 表格
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_font(hdr[i], size=10.5, bold=True, align='center', font=HEI)
        hdr[i].paragraphs[0].paragraph_format.space_before = Pt(2)
        hdr[i].paragraphs[0].paragraph_format.space_after = Pt(2)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            set_cell_font(cells[i], size=10.5, bold=False, align='center')
            cells[i].paragraphs[0].paragraph_format.space_before = Pt(1)
            cells[i].paragraphs[0].paragraph_format.space_after = Pt(1)
    if col_widths:
        for i, w in enumerate(col_widths):
            for r in t.rows:
                r.cells[i].width = Cm(w)
    # 禁止跨页
    for r in t.rows:
        tr_pr = r._tr.get_or_add_trPr()
        el = OxmlElement('w:cantSplit')
        el.set(qn('w:val'), 'true')
        tr_pr.append(el)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def add_body(doc, text, bold=False, indent=True, size=10.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if indent:
        p.paragraph_format.first_line_indent = Pt(21)  # 2字符
    r = p.add_run(text)
    r.font.name = SONG
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.element.rPr.rFonts.set(qn('w:eastAsia'), SONG)
    return p

def add_figure(doc, img_path, caption):
    doc.add_picture(img_path, width=Cm(15))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before = Pt(2)
    cp.paragraph_format.space_after = Pt(6)
    cr = cp.add_run(caption)
    cr.font.name = SONG
    cr.font.size = Pt(10.5)
    cr.font.element.rPr.rFonts.set(qn('w:eastAsia'), SONG)

def add_heading(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6 if level == 1 else 4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    r = p.add_run(text)
    sz = {1: 16, 2: 14, 3: 12}[level]
    r.font.name = HEI
    r.font.size = Pt(sz)
    r.font.bold = True
    r.font.element.rPr.rFonts.set(qn('w:eastAsia'), HEI)
    # 大纲级别
    pPr = p._p.get_or_add_pPr()
    ol = OxmlElement('w:outlineLvl')
    ol.set(qn('w:val'), str(level - 1))
    pPr.append(ol)
    return p

def add_page_break(doc):
    doc.add_page_break()

def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    fld1 = OxmlElement('w:fldChar'); fld1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve')
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld2 = OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t'); t.text = '（打开后请右键“更新域”以生成页码）'
    fld3 = OxmlElement('w:fldChar'); fld3.set(qn('w:fldCharType'), 'end')
    run._r.append(fld1); run._r.append(instr); run._r.append(fld2)
    run._r.append(t); run._r.append(fld3)

# ---------- 先生成所有图表 ----------
paths = {
    'fig1': fig_path(), 'fig2': fig_ma_periods(), 'fig3': fig_turtle_stocks(),
    'fig4': fig_turtle_heatmap(), 'fig5': fig_ml_auc(), 'fig6': fig_cs_return(),
    'fig7': fig_cs_ic(), 'fig8': fig_task7(),
}
print('图表生成完成：', len(paths))

# ---------- 新建文档 ----------
doc = Document()
# 页面 A4 + 页边距
sec = doc.sections[0]
sec.page_width = Cm(21); sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)
# 默认正文样式
normal = doc.styles['Normal']
normal.font.name = SONG
normal.font.size = Pt(10.5)
normal.element.rPr.rFonts.set(qn('w:eastAsia'), SONG)
normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
normal.paragraph_format.space_after = Pt(0)

# 页脚页码
footer = sec.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fpr = fp.add_run()
ff1 = OxmlElement('w:fldChar'); ff1.set(qn('w:fldCharType'), 'begin')
fin = OxmlElement('w:instrText'); fin.set(qn('xml:space'), 'preserve'); fin.text = 'PAGE'
ff2 = OxmlElement('w:fldChar'); ff2.set(qn('w:fldCharType'), 'end')
fpr._r.append(ff1); fpr._r.append(fin); fpr._r.append(ff2)
fpr.font.name = SONG; fpr.font.size = Pt(9); fpr.font.element.rPr.rFonts.set(qn('w:eastAsia'), SONG)

# ========================= 封面 =========================
for _ in range(3):
    doc.add_paragraph()
title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run('量化交易策略开发与机器学习应用\n综合学习报告')
tr.font.name = HEI; tr.font.size = Pt(22); tr.font.bold = True
tr.font.element.rPr.rFonts.set(qn('w:eastAsia'), HEI)
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_before = Pt(12)
sr = sub.add_run('——基于量化交易工作坊前七项任务的实践与总结')
sr.font.name = SONG; sr.font.size = Pt(13)
sr.font.element.rPr.rFonts.set(qn('w:eastAsia'), SONG)
for _ in range(4):
    doc.add_paragraph()
for line, sz in [('作者：林富强', 14), ('单位：北京大学 · 量化交易工作坊', 13),
                 ('完成日期：2026年7月24日', 13)]:
    pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pp.paragraph_format.space_after = Pt(6)
    rr = pp.add_run(line); rr.font.name = SONG; rr.font.size = Pt(sz)
    rr.font.element.rPr.rFonts.set(qn('w:eastAsia'), SONG)

add_page_break(doc)

# ========================= 目录 =========================
add_heading(doc, '目录', 1)
add_toc(doc)
add_page_break(doc)

# ========================= 摘要 =========================
add_heading(doc, '摘要', 1)
add_body(doc,
    '本报告系统总结了量化交易工作坊前七项任务的实践过程与学习成果。研究以长江电力、'
    '贵州茅台、中国平安等标的的真实行情数据为基础，依次完成了数据引擎搭建、技术指标构造、'
    '均线交叉与海龟突破两类经典策略的回测、机器学习模型在金融场景中的应用，以及基于多因子的'
    '截面选股与实盘模拟。主要结论如下：均线交叉策略在震荡市中通过降低交易频率有效抑制了回撤；'
    '海龟策略对高波动标的表现更优，但整体胜率偏低；机器学习在结构化医学数据上判别力强，'
    '而直接预测股价短期涨跌的模型判别能力普遍低于随机基准，说明线性时序规律有限；'
    '截面选股中极端梯度提升模型凭借唯一为正的信息系数跑赢等权基准。报告据此提出针对性的'
    '改进建议，并对多策略系统与深度学习等方向作出展望。')
add_page_break(doc)

# ========================= 第一章 =========================
add_heading(doc, '第一章  量化交易核心概念', 1)

add_heading(doc, '1.1  量化交易的基本概念', 2)
add_body(doc,
    '量化交易是指以数学模型和计算机程序替代人工主观判断，将投资逻辑转化为可回测、'
    '可执行的交易信号的一种投资方式。其研究对象的原始素材是市场产生的各类数据，'
    '包括逐日成交的价格与成交量，以及反映公司财务与宏观环境的基本面信息。'
    '在方法层面，量化研究通常从两条路径切入：一是技术面分析，它建立于“市场价格包容一切信息、'
    '价格沿趋势运动、历史会重演”三大假设之上，借助移动均线、相对强弱、指数平滑异同移动平均线、'
    '布林带、随机指标等工具刻画买卖力量与超买超卖状态；二是基本面分析，从宏观经济、'
    '所属行业与企业自身三个层级评估内在价值，常以估值水平与财务报表为核心依据。')
add_body(doc,
    '本工作坊的前两项任务即从这两类基础概念起步：任务一搭建了获取个股日线数据的自动化引擎，'
    '任务二对数据进行缺失值与统计分布诊断，并逐一实现了相对强弱、指数平滑异同移动平均线、'
    '布林带与随机指标四类技术指标的工程化计算。上述基础能力构成了后续所有策略研究的共同底座。')

add_heading(doc, '1.2  量化交易的核心价值', 2)
add_body(doc,
    '与依赖经验与情绪的传统人工交易相比，量化交易的核心价值可归纳为六个方面。'
    '第一，纪律性：策略由程序严格执行，规避了恐惧与贪婪对决策的干扰。'
    '第二，执行速度：算法可在极短时间完成下单，捕捉转瞬即逝的机会。'
    '第三，数据处理规模：能够同时监控成百上千只标的，超越人脑的信息负荷上限。'
    '第四，可验证性：任何策略都可在历史数据上回测，用统一指标客观评估优劣。'
    '第五，风险可控：止损、仓位管理与分散化可由规则自动落地。'
    '第六，可复用与可扩展：验证有效的策略可反复运行，并便捷地推广到新市场。'
    '正是这些特质，使量化方法成为现代投资研究的主流范式，也为本报告后续七项任务的递进式'
    '学习提供了清晰的价值主线。')

add_figure(doc, paths['fig1'], '图1  七项任务的学习路径与能力演进')

# ========================= 第二章 =========================
add_heading(doc, '第二章  量化交易策略综合分析', 1)

add_heading(doc, '2.1  均线交叉策略（任务三）', 2)
add_body(doc,
    '均线交叉策略是最具代表性的趋势类策略：当短期均线由下向上穿过长期均线时视为买入信号，'
    '反之视为卖出信号。任务三以长江电力为基准标的，将原始单体代码重构为十个职责单一的模块，'
    '系统比较了不同均线周期、四种仓位管理模式以及多只标的的表现。'
    '回测表明，在样本区间内该策略虽未取得显著绝对收益，但相对“买入并持有”均获得正超额收益，'
    '说明择时机制在震荡下行市中起到了一定的保护作用。')
add_table(doc, '表1  均线交叉策略代表参数与标的绩效',
    ['参数组合', '总收益率(%)', '年化收益率(%)', '超额收益(%)', '最大回撤(%)', '夏普比率'],
    [['均线5/10', '-1.15', '-1.30', '4.95', '-2.17', '-2.18'],
     ['均线5/20', '-0.89', '-1.00', '5.21', '-1.78', '-1.80'],
     ['均线10/20', '-0.75', '-0.85', '5.35', '-1.45', '-2.29'],
     ['均线10/30', '-0.71', '-0.83', '2.21', '-1.90', '-1.85'],
     ['长江电力', '-0.89', '-1.00', '5.21', '-1.78', '-1.80'],
     ['贵州茅台', '-0.90', '-1.02', '14.30', '-2.10', '-2.02'],
     ['中国平安', '-1.75', '-1.98', '13.73', '-1.99', '-3.61']],
    col_widths=[3.0, 2.6, 2.6, 2.6, 2.6, 2.6])
add_body(doc,
    '由表1可见，较长的均线组合（如10/30）在降低交易频率的同时回撤更小；'
    '在仓位管理上，固定金额模式以最小的回撤（-0.72%）表现最稳健。'
    '值得强调的是，三标的相对买入持有均取得正超额收益，其中贵州茅台与中国平安的超额收益'
    '分别达14.30%与13.73%，印证了择时在该区间相对“躺平”的优势。这一结论也提示我们：'
    '策略评价不能只看绝对收益，更要看相对基准的增量价值（附录建议一）。')
add_figure(doc, paths['fig2'], '图2  均线交叉策略不同参数组合的超额收益')

add_heading(doc, '2.2  海龟交易策略（任务四）', 2)
add_body(doc,
    '海龟策略是典型的突破型趋势系统，以唐奇安通道（一定周期内的最高价/最低价通道）触发入场，'
    '以真实波幅度量风险并管理仓位与止损。任务四以三只标的、八组入场/出场周期共24种参数组合进行了'
    '全网格回测，并统一采用“红涨绿跌”的本土化配色。结果显示：该策略对高波动标的更友好——'
    '在中国平安上不仅取得2.43%的正收益，更获得18.40%的超额收益，是全部组合中唯一夏普比率为正的；'
    '而在长江电力这类低波动标的上，多数参数组合跑输买入持有。')
add_table(doc, '表2  海龟策略代表结果（基准参数20/10）',
    ['标的', '总收益率(%)', '年化收益率(%)', '超额收益(%)', '最大回撤(%)', '夏普比率'],
    [['长江电力', '-6.20', '-6.95', '-0.33', '-7.41', '-1.72'],
     ['贵州茅台', '-5.47', '-6.16', '9.81', '-6.94', '-1.42'],
     ['中国平安', '2.43', '2.75', '18.40', '-9.23', '0.12']],
    col_widths=[3.0, 2.6, 2.6, 2.6, 2.6, 2.6])
add_body(doc,
    '图3与图4进一步揭示了策略表现的结构性差异：超额收益高度依赖标的选择与参数，'
    '在中国平安上几乎所有参数组合都显著跑赢基准，而在长江电力上多为小幅跑输。'
    '这一发现直接呼应了多策略配置的必要性（见2.6节，亦见附录建议六）。')
add_figure(doc, paths['fig3'], '图3  海龟策略在三只标的上的超额收益')
add_figure(doc, paths['fig4'], '图4  海龟策略24组参数的超额收益（红=跑赢，绿=跑输）')

add_heading(doc, '2.3  机器学习选股策略（任务五、任务六）', 2)
add_body(doc,
    '任务五首次将机器学习引入量化：用五种分类模型分别在结构化医学数据与长江电力股价数据上做对比。'
    '图5清楚显示，模型在医学数据上的判别能力接近满分，而在股票数据上全部跌至随机基准附近，'
    '说明“用技术指标直接预测次日涨跌”这一朴素思路面临有效市场假说的强力约束。')
add_table(doc, '表3  机器学习模型在两类数据集上的表现',
    ['数据集', '最优模型', '最高准确率', '最高AUC', '说明'],
    [['乳腺癌（结构化）', '逻辑回归', '98.25%', '0.995', '特征区分度高，模型表现优异'],
     ['长江电力（股价）', '随机森林', '46.67%', '0.46', '判别能力接近随机，预测困难']],
    col_widths=[3.2, 2.8, 2.4, 2.0, 5.6])
add_figure(doc, paths['fig5'], '图5  机器学习模型在两类数据集上的判别能力对比')
add_body(doc,
    '任务六将问题升级为“全市场截面选股”：以100只A股、12个技术因子构建季度截面面板，'
    '用四种模型预测下季度收益并选取前30只持仓、按季度调仓。结果（表4、图6、图7）表明，'
    '极端梯度提升以3.10%的年化收益跑赢2.25%的等权基准，且是唯一信息系数为正的模型（0.0133），'
    '说明其在捕捉截面非线性规律上优于线性与单树模型；而线性回归因线性假设过强表现最差。')
add_table(doc, '表4  截面选股四种模型绩效（测试期2021—2024）',
    ['模型', '年化收益(%)', '最大回撤(%)', '夏普比率', '信息系数'],
    [['线性回归', '-3.83', '-22.77', '-0.22', '-0.0869'],
     ['决策树', '2.40', '-16.92', '0.14', '-0.0260'],
     ['随机森林', '-1.30', '-20.45', '-0.07', '-0.0040'],
     ['极端梯度提升', '3.10', '-16.61', '0.17', '0.0133'],
     ['等权基准', '2.25', '—', '—', '—']],
    col_widths=[3.2, 2.6, 2.6, 2.6, 2.4])
add_figure(doc, paths['fig6'], '图6  截面选股四种模型的年化收益率')
add_figure(doc, paths['fig7'], '图7  截面选股各模型的平均信息系数')

add_heading(doc, '2.4  实盘模拟策略（任务七）', 2)
add_body(doc,
    '任务七将多因子截面选股落地到聚宽研究平台，进行回测与实盘模拟。长区间（2020年7月至2026年6月，'
    '初始资金10万元）年化收益2.97%、最大回撤24.28%、夏普比率-0.080，跑输同期沪深300的19.58%；'
    '而在样本外短区间，年化收益衰减至-8.44%、夏普比率-0.701，表现下滑明显。'
    '这再次印证了样本外衰减是量化策略的普遍挑战，也说明该研究仍处于方法验证阶段，尚未进入实盘部署。')
add_table(doc, '表5  实战策略（任务七）关键指标',
    ['区间', '年化收益(%)', '最大回撤(%)', '夏普比率', '对比基准'],
    [['长区间', '2.97', '24.28', '-0.080', '沪深300同期19.58%'],
     ['样本外', '-8.44', '—', '-0.701', '表现明显衰减']],
    col_widths=[2.4, 2.6, 2.6, 2.6, 5.8])
add_figure(doc, paths['fig8'], '图8  实战策略长区间与样本外表现对比')

add_heading(doc, '2.5  不同策略的优缺点、适用场景与市场表现比较', 2)
add_body(doc,
    '综合上述四类策略，可归纳其特性差异。均线交叉策略逻辑透明、实现简单、回撤可控，'
    '适合震荡与温和趋势市，但在单边趋势中易频繁假突破；海龟策略趋势捕捉能力强、风险预算清晰，'
    '在波动充足的标的上优势明显，但对低波动标的不敏感且交易较频繁；机器学习选股策略信息含量高、'
    '可处理非线性，适合截面排序与中低频调仓，但高度依赖特征质量，且对过拟合敏感；'
    '实盘模拟策略更贴近真实交易，但受行情阶段与成本影响大，样本外稳定性是核心瓶颈。'
    '从市场表现看，择时类策略在下行市中多以“少亏”体现价值，而选股类策略在横盘市中更易积累超额收益。')

add_heading(doc, '2.6  策略关联性与互补性：多策略系统的构建思路', 2)
add_body(doc,
    '四类策略并非彼此孤立，而是存在清晰的互补关系：均线交叉与海龟同属趋势类，可互为参数层面的'
    '分散；机器学习选股提供独立于价格趋势的截面 alpha，与择时策略低相关。由此可构建多策略系统：'
    '第一，按策略逻辑分池，趋势类与截面类各占一定权重；第二，引入市场状态识别，'
    '在趋势市提高突破类策略权重，在震荡市提高选股类权重；第三，在组合层做风险预算与仓位优化，'
    '以整体夏普比率为目标而非单一策略收益。该框架将“单点策略”升级为“策略组合”，'
    '是本报告最重要的综合结论之一，亦对应附录建议六与建议三。')

# ========================= 第三章 =========================
add_heading(doc, '第三章  机器学习在量化交易中的应用总结', 1)

add_heading(doc, '3.1  数据预处理的关键要点', 2)
add_body(doc,
    '数据预处理是模型成败的前提。实践中需把握三点：其一，金融数据具有严格的时间顺序，'
    '必须按时间划分训练与测试集，严禁随机打乱，否则会引入“未来函数”导致虚假高收益；'
    '其二，复权价应以收益率累积方式构建，规避行情接口限流与字段缺失；'
    '其三，截面面板中的缺失值需在前端统一处理（填充或剔除），否则会直接中断模型训练。'
    '任务五与任务六均严格遵循时间顺序划分，并将预处理封装为独立模块，提升了可复现性。')

add_heading(doc, '3.2  特征工程的关键要点', 2)
add_body(doc,
    '特征决定了模型上限。本报告使用的因子覆盖动量（不同周期收益率）、均线偏离（收盘价与均线的比）、'
    '波动（收益率标准差年化）、流动性（成交额与市值比）与技术面（相对强弱）五类共12个因子。'
    '核心经验是：原始技术指标对短期涨跌的预测力有限（任务五股价 AUC 低于0.5），'
    '必须引入更丰富的因子维度——如基本面、资金面、舆情与宏观变量——才能支撑有效学习（见附录建议四）。'
    '同时，特征需做截面标准化以消除量纲差异，并对极端值做合理裁剪。')

add_heading(doc, '3.3  模型选择与训练的关键要点', 2)
add_body(doc,
    '模型选择应匹配问题结构。二分类可选用逻辑回归、决策树、随机森林与近邻等可解释模型作为基线；'
    '收益率预测与排序宜用回归或梯度提升模型。训练时须固定随机种子以保证可复现，'
    '并以交叉验证而非单次切分评估稳定性。本报告在截面选股中对比了线性、树与梯度提升三类，'
    '最终极端梯度提升因正则化与集成优势取得最佳样本外表现，说明在非线性截面关系中，'
    '具备正则化的集成模型通常优于简单线性与单树模型。')

add_heading(doc, '3.4  模型评估与优化的关键要点', 2)
add_body(doc,
    '评估不能停留于准确率。对分类问题应综合混淆矩阵、精确率、召回率、F1与AUC；'
    '对选股问题应以信息系数、夏普比率、最大回撤与胜基准率为核心。优化方向包括：'
    '超参数搜索、特征筛选、模型集成与堆叠，以及最关键的是以样本外与滚动窗口验证抑制过拟合'
    '（附录建议一）。任务六中极端梯度提升是唯一信息系数为正的模型，正是评估视角从“准确率”'
    '转向“排序能力”后得到的更有意义结论。')

add_heading(doc, '3.5  优势、局限与未来趋势', 2)
add_body(doc,
    '机器学习的核心优势在于自动挖掘非线性与高维规律、处理海量数据并持续迭代；'
    '其局限同样突出：金融数据信噪比低、分布易变，模型易过拟合且可解释性弱，'
    '直接预测价格短期方向往往失效。展望未来，三个方向值得投入：一是引入深度学习'
    '（序列模型与时序注意力机制）刻画价格动态；二是以“因子+模型+组合”三段式框架'
    '替代端到端黑箱；三是将机器学习输出作为多策略系统中的一个 alpha 源，而非孤立决策。'
    '（相关路线见附录建议四、建议五。）')

# ========================= 第四章 =========================
add_heading(doc, '第四章  结论与展望', 1)

add_heading(doc, '4.1  主要收获与体会', 2)
add_body(doc,
    '通过前七项任务的递进式学习，本人在三方面获得实质性提升。其一是认知层面：'
    '从“量化即写公式”深化为对策略逻辑、风险来源与评估体系的系统理解，认识到'
    '“相对基准的超额收益”比绝对收益更能反映策略价值。其二是技术层面：'
    '掌握了从数据获取、指标构造、回测引擎、机器学习建模到平台部署的完整工具链，'
    '并养成了模块化、可复现的工程习惯（任务三的重构、任务四的配色统一即为例证）。'
    '其三是实践层面：亲历了样本外衰减、过拟合与执行成本等真实挑战，'
    '体会到量化交易“知易行难”的严谨性。')

add_heading(doc, '4.2  未来学习与研究方向展望', 2)
add_body(doc,
    '面向未来，研究将沿以下路径深化。第一，补强验证方法：以滚动窗口与样本外检验'
    '替代单次回测，从源头抑制过拟合（附录建议一）。第二，拓展数据维度：'
    '纳入基本面、资金面与另类数据，突破单一技术指标的瓶颈（附录建议四）。'
    '第三，完善组合与风控：在选股之外引入仓位优化、止损与行业中性，'
    '降低尾部风险（附录建议三）。第四，探索先进模型：尝试序列模型与模型堆叠，'
    '捕捉非线性时序（附录建议五）。第五，构建多策略配置框架：按市场状态动态分配权重，'
    '实现策略间互补（附录建议六）。第六，做实执行层：在模拟盘中计入滑点与手续费压力测试，'
    '为真正实盘打下基础（附录建议七）。通过上述努力，期望由“会做策略”走向“做对策略”。')

# ========================= 附录 =========================
add_heading(doc, '附录  针对性改进建议', 1)
add_body(doc,
    '以下建议基于前七项任务的复盘与导师反馈整理，正文中已按编号引用。',
    indent=False)
appendix = [
    ('建议一（验证方法）', '以滚动窗口回测与严格样本外检验替代单次历史回测，'
     '对多参数策略（如任务三、任务四）增加参数稳定性分析，从源头抑制过拟合与样本内最优陷阱。'),
    ('建议二（数据广度）', '将回测标的由三只个股扩展为更宽泛的股票池与更长历史区间，'
     '降低单一标的与特定行情阶段带来的结论偏差（任务七样本外大幅衰减即为例证）。'),
    ('建议三（组合与风控）', '在选股信号之外引入组合层优化：仓位管理、止损纪律、'
     '行业中性与最大回撤约束，使策略从“选得准”升级为“控得住”（任务六回撤仍达-16%）。'),
    ('建议四（特征升级）', '突破单一技术指标，纳入基本面、资金面、估值与另类（舆情/宏观）因子，'
     '提升特征信息含量，直指任务五股价预测判别力不足的根因。'),
    ('建议五（模型进阶）', '在树模型基础上尝试序列模型与模型堆叠，'
     '捕捉价格序列中的非线性与长期依赖，提升截面与时序预测能力。'),
    ('建议六（多策略配置）', '构建趋势类与截面类并行的多策略框架，'
     '按市场状态动态分配权重，以整体夏普比率为目标实现策略互补（呼应第二章2.6节）。'),
    ('建议七（执行落地）', '在实盘模拟中完整计入滑点、手续费与冲击成本，'
     '开展压力测试与模拟盘长跑，缩小研究与实盘的鸿沟（任务七尚未充分实盘）。'),
    ('建议八（工程规范）', '统一数据、指标与评价的回测框架与代码规范，'
     '沉淀可复用的模块库，提升后续研究的效率与可复现性（延续任务三模块化、任务四配色统一经验）。'),
]
for title_a, body_a in appendix:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.first_line_indent = Pt(21)
    r1 = p.add_run(title_a + '：')
    r1.font.name = SONG; r1.font.size = Pt(10.5); r1.font.bold = True
    r1.font.element.rPr.rFonts.set(qn('w:eastAsia'), SONG)
    r2 = p.add_run(body_a)
    r2.font.name = SONG; r2.font.size = Pt(10.5); r2.font.bold = False
    r2.font.element.rPr.rFonts.set(qn('w:eastAsia'), SONG)

# ---------- 保存 docx ----------
doc.save(DOCX)
print('DOCX 已保存：', DOCX)

# ---------- 转 PDF（Word COM） ----------
import win32com.client as win32
word = win32.Dispatch('Word.Application')
word.Visible = False
d = word.Documents.Open(DOCX)
try:
    d.TablesOfContents(1).Update()
except Exception as e:
    print('TOC 更新跳过：', e)
d.Fields.Update()
d.SaveAs2(PDF, FileFormat=17)
d.Close()
word.Quit()
print('PDF 已保存：', PDF)
